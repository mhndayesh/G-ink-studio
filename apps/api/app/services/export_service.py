from __future__ import annotations

"""Export rendering: turns the six story JSON files into text/markdown/docx and
ZIP bundles. Pure logic — no FastAPI types. Driven by app/api/v1/export.py.

This module is large; it is a candidate for further splitting (story / visuals /
scenes / validation renderers) — see docs/REPO-CRITIQUE.md.
"""

import io
import json
import zipfile

from app.repositories.sqlite_registry import SQLiteRegistry
from app.services.visual_prompt import (
    STYLE_PREFIX,
    canonical_camera_shot,
    compile_visual_prompt,
    negative_prompt,
    render_mode_for_cast,
    sanitize_visual_prompt,
)


def story_safe_title(files: dict, fallback: str = "story") -> tuple[str, str]:
    """Return (display_title, filesystem_safe_title) from the master_story file."""
    ms = files.get("master_story", {}) or {}
    title = _safe(ms.get("title") or ms.get("story_title"), fallback)
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip().replace(" ", "_")
    return title, safe


# ─── helpers ────────────────────────────────────────────────────────────────

def _get_all_files(story_id: str, registry: SQLiteRegistry) -> dict:
    out: dict = {}
    for ft in ["master_story", "characters", "plot_outline", "plot_workspace", "chapter_script"]:
        rec = registry.get_current_file(story_id, ft)
        if rec:
            out[ft] = rec.get("json_copy", {})
    return out


def _safe(v: object, fallback: str = "") -> str:
    if v is None:
        return fallback
    return str(v).strip() or fallback


def _selected(v: object) -> str:
    if isinstance(v, dict):
        selected = v.get("selected")
        if isinstance(selected, list):
            return ", ".join(_text_list(selected))
        return _safe(selected)
    if isinstance(v, list):
        return ", ".join(_text_list(v))
    return _safe(v)


def _as_list(v: object) -> list:
    return v if isinstance(v, list) else []


def _text_list(v: object) -> list[str]:
    out: list[str] = []
    for item in _as_list(v):
        if isinstance(item, str):
            text = _safe(item)
        elif isinstance(item, dict):
            text = _safe(
                item.get("name")
                or item.get("character_name")
                or item.get("profile_name")
                or item.get("title")
                or item.get("label")
                or item.get("summary")
                or item.get("description")
                or item.get("relationship_id")
                or item.get("threat_id_or_name")
            )
        else:
            text = _safe(item)
        if text:
            out.append(text)
    return out


def _append_field(lines: list[str], label: str, value: object, indent: str = "") -> None:
    if isinstance(value, list):
        text = ", ".join(_text_list(value))
    else:
        text = _safe(value)
    if text:
        lines.append(f"{indent}{label}: {text}")


def _story_title(ms: dict, fallback: str = "Untitled Story") -> str:
    foundation = ms.get("story_foundation") if isinstance(ms.get("story_foundation"), dict) else {}
    return _safe(ms.get("title") or ms.get("story_title") or foundation.get("story_title"), fallback)


def _chapter_header(chapter: dict) -> str:
    ch_num = _safe(chapter.get("chapter_number") or chapter.get("chapter_id"), "?")
    ch_title = _safe(chapter.get("chapter_title") or chapter.get("title"))
    return f"Chapter {ch_num}{': ' + ch_title if ch_title else ''}"


def _group_scenes_by_chapter(scene_cards: list[dict]) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = {}
    for scene in scene_cards:
        if not isinstance(scene, dict):
            continue
        grouped.setdefault(_safe(scene.get("chapter_id")), []).append(scene)
    for scenes in grouped.values():
        scenes.sort(key=lambda s: int(s.get("scene_order") or 0))
    return grouped


def _panel_story_lines(pages: list[dict]) -> list[str]:
    lines: list[str] = []
    for page in pages:
        if not isinstance(page, dict):
            continue
        page_label = _safe(page.get("page_number") or page.get("page_id"))
        page_bits = []
        if _safe(page.get("page_purpose")):
            page_bits.append(_safe(page.get("page_purpose")))
        if _safe(page.get("page_mood")):
            page_bits.append(f"Mood: {_safe(page.get('page_mood'))}")
        if page_bits:
            lines.append(f"Page {page_label}: " + " ".join(page_bits))
        for panel in _as_list(page.get("panels")):
            if not isinstance(panel, dict):
                continue
            bits = [
                _safe(panel.get("visual")),
                _safe(panel.get("character_action")),
                _safe(panel.get("narration")),
            ]
            for dialogue in _as_list(panel.get("dialogue")):
                if not isinstance(dialogue, dict):
                    continue
                text = _safe(dialogue.get("text"))
                if not text:
                    continue
                speaker = _safe(dialogue.get("speaker_name") or dialogue.get("speaker") or dialogue.get("speaker_id"))
                bits.append(f"{speaker}: {text}" if speaker else text)
            panel_text = " ".join(bit for bit in bits if bit)
            if panel_text:
                panel_no = _safe(panel.get("panel_number") or panel.get("panel_id"))
                lines.append(f"Panel {panel_no}: {panel_text}")
    return lines


# ─── story assembly ──────────────────────────────────────────────────────────

def _lines_to_text(lines: list[str]) -> str:
    return "\n".join(lines) + "\n"


def _lines_to_markdown(lines: list[str]) -> str:
    """Convert the assembled lines to Markdown with proper heading levels."""
    md: list[str] = []
    for i, line in enumerate(lines):
        # Detect heading markers based on next-line underlines
        if i + 1 < len(lines):
            nxt = lines[i + 1]
            if nxt and all(c == "=" for c in nxt) and len(nxt) >= 3:
                md.append(f"# {line}")
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= 3:
                md.append(f"## {line}")
                continue
        # Skip the underline rows themselves
        if line and (all(c == "=" for c in line) or all(c == "-" for c in line)):
            md.append("")
            continue
        md.append(line)
    return "\n".join(md) + "\n"


def _lines_to_docx(lines: list[str], title: str) -> bytes:
    """Generate a .docx file from assembled lines. Falls back to plain bytes on import error."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[attr-defined]

        doc = Document()

        # Remove default empty paragraph
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        i = 0
        while i < len(lines):
            line = lines[i]
            nxt = lines[i + 1] if i + 1 < len(lines) else ""

            if nxt and all(c == "=" for c in nxt) and len(nxt) >= 3:
                doc.add_heading(line, level=1)
                i += 2  # skip underline
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= 3:
                doc.add_heading(line, level=2)
                i += 2
                continue
            if line and (all(c == "=" for c in line) or all(c == "-" for c in line)):
                i += 1
                continue
            if line == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # python-docx not installed — return plain text bytes with .docx extension
        return _lines_to_text(lines).encode("utf-8")


# ─── visuals assembly ────────────────────────────────────────────────────────

def _all_chapter_scripts(story_id: str, registry: SQLiteRegistry) -> list[dict]:
    """Walk every version's chapter_script.json snapshot, deduped by chapter_id.

    Latest version per chapter wins (so an updated chapter overrides older snapshots).
    """
    rows = registry.list_files_across_versions(story_id, "chapter_script")
    by_chapter: dict[str, dict] = {}
    for row in rows:
        data = row.get("json_copy", {}) or {}
        if not _as_list(data.get("pages")):
            continue
        ch_id = _safe((data.get("chapter_metadata") or {}).get("chapter_id")) or _safe(row.get("version_id"))
        if not ch_id:
            continue
        # Latest wins (rows are ordered by version_number ASC).
        by_chapter[ch_id] = data
    # Sort by chapter_number
    return sorted(
        by_chapter.values(),
        key=lambda d: int((d.get("chapter_metadata") or {}).get("chapter_number") or 0),
    )


def _appearance_block(profile: dict) -> dict:
    block = profile.get("appearance_and_visual_design", {})
    if isinstance(block, dict):
        details = block.get("appearance_details", {})
        if isinstance(details, dict) and details:
            return details
    # Fallback: some profiles may flatten the fields onto the root.
    return profile if any(profile.get(k) for k in ("hair_style", "main_outfit_description", "color_palette")) else {}


def _character_reference_lines(characters: dict) -> list[str]:
    """Per-character visual reference sheet for the artist."""
    lines: list[str] = []
    profiles = (characters.get("created_major_character_profiles", []) or []) + (characters.get("created_side_character_profiles", []) or [])
    if not profiles:
        return lines
    lines += ["CHARACTER REFERENCE SHEETS", "=" * 26, ""]
    for profile in profiles:
        if not isinstance(profile, dict):
            continue
        name = _safe(profile.get("character_name"))
        if not name:
            continue
        # Spec: character names lowercase — parser silently drops fields otherwise
        name = name.lower()
        role = _profile_role(profile)
        header = name + (f" - {role}" if role else "")
        lines += [header, "-" * len(header)]
        details = _appearance_block(profile)
        for key, label in [
            ("age_range", "Age"),
            ("gender_presentation", "Gender presentation"),
            ("height", "Height"),
            ("body_type", "Body type"),
            ("silhouette_shape", "Silhouette"),
            ("face_shape", "Face shape"),
            ("skin_tone_or_markings", "Skin / markings"),
            ("hair_style", "Hair style"),
            ("hair_color", "Hair color"),
            ("eye_shape", "Eye shape"),
            ("eye_color", "Eye color"),
            ("clothing_style", "Clothing style"),
            ("main_outfit_description", "Main outfit"),
            ("iconic_item", "Iconic item"),
            ("expression_style", "Expression style"),
            ("pose_language", "Pose language"),
            ("manga_panel_presence", "Panel presence"),
            ("first_impression_visual", "First impression"),
        ]:
            v = details.get(key)
            text = ", ".join(_text_list(v)) if isinstance(v, list) else _safe(v)
            if text:
                lines.append(f"  {label}: {text}")
        for key, label in [
            ("distinctive_features", "Distinctive features"),
            ("scars_or_birthmarks", "Scars / birthmarks"),
            ("alternate_outfits", "Alternate outfits"),
            ("accessories", "Accessories"),
            ("weapons_or_tools_visible", "Weapons / tools"),
            ("color_palette", "Color palette"),
        ]:
            text = ", ".join(_text_list(details.get(key, [])))
            if text:
                lines.append(f"  {label}: {text}")
        symbol = _safe(details.get("visual_symbol_or_motif"))
        if symbol:
            lines.append(f"  Visual symbol / motif: {symbol}")
        contrast = _safe(details.get("visual_contrast_with_other_characters"))
        if contrast:
            lines.append(f"  Visual contrast: {contrast}")
        lines.append(f"  AI prompt (positive): {compile_visual_prompt(', '.join(_character_visual_phrases(details, profile)))}")
        lines.append(f"  AI prompt (negative): {negative_prompt(details.get('negative_prompt_notes'))}")
        lines.append("")
    return lines


def _location_index_lines(scripts: list[dict], loc_by_id: dict[str, dict] | None = None) -> list[str]:
    """Cross-chapter index of every location and which panels reference it.

    Resolves location names from (in order of preference):
      1. page.location_id → loc_by_id lookup (set by Visuals Studio)
      2. scene_breakdown.location_id → loc_by_id lookup
      3. scene_breakdown.location string
      4. page.location string (legacy)
    """
    locs: dict[str, list[str]] = {}
    loc_map = loc_by_id or {}
    for cs in scripts:
        ch_meta = cs.get("chapter_metadata", {}) or {}
        ch_num = _safe(ch_meta.get("chapter_number") or ch_meta.get("chapter_id"), "?")
        scene_breakdown = _as_list(cs.get("chapter_scene_breakdown"))
        scene_loc_id: dict[str, str] = {}   # scene_id → location_id
        scene_loc_name: dict[str, str] = {} # scene_id → location name string
        for s in scene_breakdown:
            if not isinstance(s, dict):
                continue
            sid = _safe(s.get("scene_id"))
            if s.get("location_id"):
                scene_loc_id[sid] = _safe(s.get("location_id"))
            if s.get("location"):
                scene_loc_name[sid] = _safe(s.get("location"))
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            scene_id = _safe(page.get("scene_id"))
            page_num = _safe(page.get("page_number") or page.get("page_id"))
            ref = f"Ch.{ch_num} Pg.{page_num}"
            # Priority 1: page.location_id (set by Visuals Studio)
            page_lid = _safe(page.get("location_id"))
            if page_lid and page_lid in loc_map:
                loc_name = _safe(loc_map[page_lid].get("name"))
                if loc_name:
                    locs.setdefault(loc_name, []).append(ref)
                    continue
            # Priority 2: scene_breakdown.location_id → name
            scene_lid = scene_loc_id.get(scene_id, "")
            if scene_lid and scene_lid in loc_map:
                loc_name = _safe(loc_map[scene_lid].get("name"))
                if loc_name:
                    locs.setdefault(loc_name, []).append(ref)
                    continue
            # Priority 3: scene_breakdown.location string
            loc_str = scene_loc_name.get(scene_id, "") or _safe(page.get("location"))
            if loc_str:
                locs.setdefault(loc_str, []).append(ref)
    if not locs:
        return []
    lines = ["LOCATION INDEX", "=" * 14, ""]
    for loc in sorted(locs.keys()):
        lines.append(f"{loc}")
        lines.append(f"  Used in: {', '.join(locs[loc])}")
    lines.append("")
    return lines


def _panel_named_cast_count(panel: dict) -> int:
    """Distinct named speaking characters in a panel (excludes Narrator / extras).

    Best signal available at export time for the render-mode policy; under-counts
    silent characters, which is fine — it never over-promises a 2-ref stitch.
    """
    names: set[str] = set()
    for d in _as_list(panel.get("dialogue", [])):
        if not isinstance(d, dict):
            continue
        spk = _safe(d.get("speaker_name") or d.get("speaker") or d.get("speaker_id")).strip().lower()
        if spk and spk not in ("narrator", "narration", "?", "background character", "background", "sfx"):
            names.add(spk)
    return len(names)


def _panel_full_block(panel: dict, indent: str = "    ") -> list[str]:
    """Render every production-relevant panel field for the artist."""
    out: list[str] = []
    panel_num = _safe(panel.get("panel_number") or panel.get("panel_id"))
    shot_block = panel.get("camera_shot")
    camera_raw = _selected(shot_block) if isinstance(shot_block, dict) else _safe(shot_block)
    camera = canonical_camera_shot(camera_raw)  # BUNDLE-AUDIT #2: drop "Action Shot"/"Reaction Shot" etc.
    size_block = panel.get("panel_size")
    size = _selected(size_block) if isinstance(size_block, dict) else _safe(size_block)
    pacing_block = panel.get("pacing")
    pacing = _selected(pacing_block) if isinstance(pacing_block, dict) else _safe(pacing_block)
    tags = [t for t in [size, camera, pacing] if t]
    header = f"{indent}Panel {panel_num}"
    if tags:
        header += f" [{' / '.join(tags)}]"
    out.append(header)
    for key, label in [
        ("visual", "Visual"),
        ("character_action", "Action"),
        ("background_details", "Background"),
        ("facial_expression", "Expression"),
        ("pose_or_body_language", "Pose"),
        ("mood", "Mood"),
        ("narration", "Narration"),
        ("continuity_notes", "Continuity"),
        ("custom_panel_details", "Custom"),
    ]:
        v = _safe(panel.get(key))
        if not v:
            continue
        # BUNDLE-AUDIT #8: an "N/A …" expression on an object-only panel is noise.
        if key == "facial_expression" and v.strip().lower().startswith(("n/a", "none", "not applicable")):
            continue
        out.append(f"{indent}  {label}: {v}")
    sfx_items = _as_list(panel.get("sound_effects"))
    sfx_parts = []
    for s in sfx_items:
        if isinstance(s, dict):
            t = _safe(s.get("sfx_text"))
            meaning = _safe(s.get("sfx_meaning"))
            if t:
                sfx_parts.append(f"{t}{f' ({meaning})' if meaning else ''}")
        elif isinstance(s, str) and s.strip():
            sfx_parts.append(s.strip())
    if sfx_parts:
        out.append(f"{indent}  SFX: {', '.join(sfx_parts)}")
    # Dialogue lines — format: Speaker (BubbleType): "text"
    for d in _as_list(panel.get("dialogue", [])):
        if not isinstance(d, dict):
            continue
        text = _safe(d.get("text"))
        if not text:
            continue
        speaker = _safe(d.get("speaker_name") or d.get("speaker") or d.get("speaker_id")) or "?"
        bubble_block = d.get("speech_bubble_type")
        bubble = _selected(bubble_block) if isinstance(bubble_block, dict) else _safe(bubble_block)
        # Skip Narrator entries — Narration: field already covers caption text (avoids duplicate lettering)
        if speaker.lower() in ("narrator", "narration") or (bubble and bubble.lower() == "narration"):
            continue
        bubble_tag = f" ({bubble})" if bubble and bubble.lower() not in ("", "normal", "standard") else ""
        out.append(f'{indent}  Dialogue: {speaker}{bubble_tag}: "{text}"')
    # Render mode — BUNDLE-AUDIT #3: derive from the in-frame named cast, not from
    # panel position. 0 chars → t2i, 1 → i2i, 2 → i2i-2refs, 3+ → i2i-2refs + a
    # warning (a 2-reference stitch can't hold a third character). The studio still
    # resolves the *effective* mode against the reference images actually on file.
    rm, rm_warn = render_mode_for_cast(_panel_named_cast_count(panel))
    out.append(f"{indent}  Render mode: {rm}" + (f"  (note: {rm_warn})" if rm_warn else ""))
    return out


def _locations_section_lines(po: dict) -> list[str]:
    """Build the LOCATIONS section from plot_outline.locations.locations[]."""
    locs_block = (po.get("locations") or {})
    locs = _as_list(locs_block.get("locations", []))
    if not locs:
        return []
    lines: list[str] = ["LOCATIONS", "=" * 9, ""]
    for loc in locs:
        if not isinstance(loc, dict):
            continue
        name = _safe(loc.get("name"))
        if not name:
            continue
        loc_type = _safe(loc.get("type"))
        header = name + (f" - {loc_type}" if loc_type else "")
        lines += [header, "-" * len(header)]
        desc = _safe(loc.get("description"))
        if desc:
            lines.append(f"  Description: {desc}")
        # Compile a clean, B&W-manga prompt: structural fragments from the prose
        # description + the LLM positive_prompt, colour/lighting/style stripped,
        # STYLE_PREFIX prepended (BUNDLE-AUDIT #6).
        lines.append(f"  AI prompt (positive): {compile_visual_prompt(loc_type, loc.get('positive_prompt'), desc)}")
        lines.append(f"  AI prompt (negative): {negative_prompt(loc.get('negative_prompt') or 'people, characters')}")
        lines.append("")
    return lines


def _build_loc_by_id(po: dict) -> dict[str, dict]:
    """Map location_id → location dict from plot_outline."""
    locs_block = (po.get("locations") or {})
    out: dict[str, dict] = {}
    for loc in _as_list(locs_block.get("locations", [])):
        if isinstance(loc, dict):
            lid = _safe(loc.get("location_id"))
            if lid:
                out[lid] = loc
    return out


def _assemble_visuals_lines(files: dict, all_scripts: list[dict] | None = None) -> list[str]:
    """Build the visual reference document.

    Aggregates panel content across ALL chapter_script snapshots (one per chapter)
    plus a character reference sheet, a LOCATIONS section, and a location index.
    """
    lines: list[str] = []
    po = files.get("plot_outline", {}) or {}
    title = _story_title(files.get("master_story", {}) or {}, "Story")
    lines += [f"VISUAL REFERENCE - {title}", "=" * (19 + len(title)), ""]

    # Character reference sheets
    lines.extend(_character_reference_lines(files.get("characters", {}) or {}))

    # LOCATIONS section — before chapters, so studio can parse location prompts
    lines.extend(_locations_section_lines(po))

    # Choose source: cross-version aggregate or single current chapter_script
    scripts: list[dict]
    if all_scripts is not None:
        scripts = all_scripts
    else:
        cs = files.get("chapter_script", {}) or {}
        scripts = [cs] if _as_list(cs.get("pages")) else []

    if not scripts:
        lines.append("No script pages found. Generate the Manga Script for at least one chapter first.")
        return lines

    # Build location_id → name lookup (shared by index + per-page headers)
    loc_by_id = _build_loc_by_id(po)

    # Location index — cross-chapter usage map
    lines.extend(_location_index_lines(scripts, loc_by_id=loc_by_id))

    # Per-chapter panel breakdown
    lines += ["CHAPTERS", "=" * 8, ""]
    for cs in scripts:
        metadata = cs.get("chapter_metadata", {}) or {}
        ch_num = _safe(metadata.get("chapter_number") or metadata.get("chapter_id"), "?")
        ch_title = _safe(metadata.get("chapter_title"))
        ch_status = _safe(metadata.get("chapter_status"))
        ch_header = f"Chapter {ch_num}{': ' + ch_title if ch_title else ''}"
        lines += [ch_header, "-" * len(ch_header)]
        if ch_status:
            lines.append(f"Status: {ch_status}")
        scene_breakdown = _as_list(cs.get("chapter_scene_breakdown"))
        scene_titles: dict[str, str] = {}
        scene_loc_id: dict[str, str] = {}   # scene_id → location_id from breakdown
        scene_loc_name: dict[str, str] = {} # scene_id → location string name from breakdown
        for s in scene_breakdown:
            if not isinstance(s, dict):
                continue
            sid = _safe(s.get("scene_id"))
            scene_titles[sid] = _safe(s.get("scene_title") or s.get("location"))
            if s.get("location_id"):
                scene_loc_id[sid] = _safe(s.get("location_id"))
            if s.get("location"):
                scene_loc_name[sid] = _safe(s.get("location"))
        # All known location names for this chapter — used to repair a page header
        # whose stored Location doesn't match what the panels actually depict
        # (BUNDLE-AUDIT #1). Names shorter than 4 chars are skipped (false-match risk).
        known_loc_names = {
            n for n in (
                [_safe(v.get("name")) for v in loc_by_id.values()]
                + list(scene_loc_name.values())
                + [_safe(l.get("name")) for l in _as_list((po.get("locations") or {}).get("locations", []))]
            ) if len(n) >= 4
        }
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            page_num = _safe(page.get("page_number") or page.get("page_id"))
            scene_id = _safe(page.get("scene_id"))
            scene_label = scene_titles.get(scene_id, scene_id)
            page_purpose = _safe(page.get("page_purpose"))
            page_mood = _safe(page.get("page_mood"))

            # Resolve location name: page.location_id → loc_by_id → name
            # Fallback: scene_breakdown location_id → name, then scene_breakdown location string
            loc_name = ""
            page_loc_id = _safe(page.get("location_id"))
            if page_loc_id and page_loc_id in loc_by_id:
                loc_name = _safe(loc_by_id[page_loc_id].get("name"))
            if not loc_name:
                scene_lid = scene_loc_id.get(scene_id, "")
                if scene_lid and scene_lid in loc_by_id:
                    loc_name = _safe(loc_by_id[scene_lid].get("name"))
            if not loc_name:
                loc_name = scene_loc_name.get(scene_id, "")

            # BUNDLE-AUDIT #1: if the stored Location isn't mentioned anywhere in
            # this page's panels but exactly one *other* known location is, the
            # upstream scene→location map is wrong — prefer what the panels show.
            page_text = " ".join(
                f"{_safe(pn.get('visual'))} {_safe(pn.get('background_details'))} {_safe(pn.get('character_action'))}"
                for pn in _as_list(page.get("panels")) if isinstance(pn, dict)
            ).lower()
            if page_text:
                in_panels = [n for n in known_loc_names if n.lower() in page_text]
                if loc_name and loc_name.lower() not in page_text and len(in_panels) == 1 and in_panels[0] != loc_name:
                    loc_name = f"{in_panels[0]} (corrected from {loc_name!r} — panels depict this setting)"
                elif not loc_name and len(in_panels) == 1:
                    loc_name = in_panels[0]

            header = f"\n  Page {page_num}"
            if scene_label:
                header += f" - Scene: {scene_label}"
            if loc_name:
                header += f" - Location: {loc_name}"
            lines.append(header)
            if page_purpose:
                lines.append(f"    Purpose: {page_purpose}")
            if page_mood:
                lines.append(f"    Mood: {page_mood}")
            for panel in _as_list(page.get("panels")):
                if isinstance(panel, dict):
                    lines.extend(_panel_full_block(panel, indent="    "))
        lines.append("")

    return lines


def _panels_csv(scripts: list[dict]) -> str:
    """Per-panel CSV for spreadsheet workflows. One row per panel across all chapters."""
    import csv
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "chapter_number", "chapter_id", "chapter_title",
        "page_number", "page_id", "scene_id",
        "panel_number", "panel_id", "panel_size", "camera_shot", "pacing",
        "visual", "character_action", "background_details",
        "facial_expression", "pose_or_body_language", "mood",
        "narration", "dialogue_count", "sfx_text",
        "continuity_notes",
    ])
    for cs in scripts:
        meta = cs.get("chapter_metadata", {}) or {}
        ch_num = _safe(meta.get("chapter_number"))
        ch_id = _safe(meta.get("chapter_id"))
        ch_title = _safe(meta.get("chapter_title"))
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            page_num = _safe(page.get("page_number"))
            page_id = _safe(page.get("page_id"))
            scene_id = _safe(page.get("scene_id"))
            for panel in _as_list(page.get("panels")):
                if not isinstance(panel, dict):
                    continue
                shot_block = panel.get("camera_shot")
                size_block = panel.get("panel_size")
                pacing_block = panel.get("pacing")
                sfx_items = _as_list(panel.get("sound_effects"))
                sfx_text = "; ".join(
                    _safe(s.get("sfx_text")) if isinstance(s, dict) else _safe(s)
                    for s in sfx_items
                    if (isinstance(s, dict) and _safe(s.get("sfx_text"))) or (isinstance(s, str) and s.strip())
                )
                shot_raw = _selected(shot_block) if isinstance(shot_block, dict) else _safe(shot_block)
                expr = _safe(panel.get("facial_expression"))
                if expr.strip().lower().startswith(("n/a", "none", "not applicable")):
                    expr = ""
                writer.writerow([
                    ch_num, ch_id, ch_title,
                    page_num, page_id, scene_id,
                    _safe(panel.get("panel_number")),
                    _safe(panel.get("panel_id")),
                    _selected(size_block) if isinstance(size_block, dict) else _safe(size_block),
                    canonical_camera_shot(shot_raw),  # BUNDLE-AUDIT #2
                    _selected(pacing_block) if isinstance(pacing_block, dict) else _safe(pacing_block),
                    _safe(panel.get("visual")),
                    _safe(panel.get("character_action")),
                    _safe(panel.get("background_details")),
                    expr,
                    _safe(panel.get("pose_or_body_language")),
                    _safe(panel.get("mood")),
                    _safe(panel.get("narration")),
                    len(_as_list(panel.get("dialogue"))),
                    sfx_text,
                    _safe(panel.get("continuity_notes")),
                ])
    return buf.getvalue()


def _character_visual_phrases(details: dict, profile: dict | None = None) -> list[str]:
    """Visual-only descriptor fragments for a character, drawn from the structured
    appearance fields (no scene / pose / lighting / colour — see visual_prompt.py)."""
    parts: list[object] = []
    for key in ("age_range", "gender_presentation", "height", "body_type", "silhouette_shape",
                "face_shape", "skin_tone_or_markings", "hair_style", "eye_shape",
                "clothing_style", "main_outfit_description", "iconic_item",
                "visual_symbol_or_motif"):
        parts.append(details.get(key))
    for key in ("distinctive_features", "scars_or_birthmarks", "accessories", "weapons_or_tools_visible"):
        parts.extend(_text_list(details.get(key, [])))
    # the LLM-written notes go last so the structured fields anchor the prompt
    parts.append(details.get("ai_image_prompt_notes"))
    phrases: list[str] = []
    seen: set[str] = set()
    for part in parts:
        for ph in sanitize_visual_prompt(part).split(", "):
            k = ph.lower()
            if ph and k not in seen:
                seen.add(k)
                phrases.append(ph)
    return phrases


def _ai_prompt_files(characters: dict) -> dict[str, str]:
    """One file per character holding a clean, B&W-manga AI image prompt (positive + negative)."""
    out: dict[str, str] = {}
    profiles = (characters.get("created_major_character_profiles", []) or []) + (characters.get("created_side_character_profiles", []) or [])
    for p in profiles:
        if not isinstance(p, dict):
            continue
        name = _safe(p.get("character_name"))
        if not name:
            continue
        details = _appearance_block(p)
        phrases = _character_visual_phrases(details, p)
        pos = compile_visual_prompt(", ".join(phrases))
        neg = negative_prompt(details.get("negative_prompt_notes"))
        if pos == STYLE_PREFIX and not phrases:
            continue
        safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in name).strip().replace(" ", "_") or "character"
        out[f"{safe_name}.txt"] = f"# Positive prompt\n{pos}\n\n# Negative prompt\n{neg}\n"
    return out


def _character_sheet_files(characters: dict) -> dict[str, str]:
    """One Markdown file per character with their full visual sheet."""
    out: dict[str, str] = {}
    profiles = (characters.get("created_major_character_profiles", []) or []) + (characters.get("created_side_character_profiles", []) or [])
    for p in profiles:
        if not isinstance(p, dict):
            continue
        name = _safe(p.get("character_name"))
        if not name:
            continue
        details = _appearance_block(p)
        if not details:
            continue
        safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in name).strip().replace(" ", "_") or "character"
        md: list[str] = [f"# {name}"]
        role = _profile_role(p)
        if role:
            md.append(f"_{role}_\n")
        for key, label in [
            ("age_range", "Age"),
            ("body_type", "Body type"),
            ("face_shape", "Face shape"),
            ("hair_style", "Hair"),
            ("hair_color", "Hair color"),
            ("eye_color", "Eye color"),
            ("clothing_style", "Clothing"),
            ("main_outfit_description", "Outfit"),
            ("iconic_item", "Iconic item"),
            ("expression_style", "Expression style"),
        ]:
            v = _safe(details.get(key))
            if v:
                md.append(f"**{label}**: {v}")
        for key, label in [
            ("distinctive_features", "Distinctive features"),
            ("accessories", "Accessories"),
            ("weapons_or_tools_visible", "Weapons / tools"),
            ("color_palette", "Color palette"),
        ]:
            text = ", ".join(_text_list(details.get(key, [])))
            if text:
                md.append(f"**{label}**: {text}")
        md.append(f"\n## AI prompt (positive)\n{compile_visual_prompt(', '.join(_character_visual_phrases(details, p)))}")
        md.append(f"\n## AI prompt (negative)\n{negative_prompt(details.get('negative_prompt_notes'))}")
        out[f"{safe_name}.md"] = "\n\n".join(md) + "\n"
    return out


# ─── routes ──────────────────────────────────────────────────────────────────

def _profile_role(profile: dict) -> str:
    """Resolve a character profile's role from the canonical template path."""
    role_block = profile.get("character_role_level")
    if isinstance(role_block, dict):
        sel = _safe(role_block.get("selected"))
        if sel:
            return sel
        custom = _safe(role_block.get("custom_character_role_level"))
        if custom:
            return custom
    return _safe(profile.get("profile_label"))


def _profile_bio(profile: dict) -> str:
    """Build a one-paragraph bio from real character template fields."""
    bits: list[str] = []
    backstory = profile.get("character_backstory_mental_state_and_community_place", {}) or {}
    bd = backstory.get("backstory_details", {}) if isinstance(backstory, dict) else {}
    if isinstance(bd, dict):
        for key in ("childhood_summary", "important_past_event", "past_trauma_or_wound"):
            v = _safe(bd.get(key))
            if v:
                bits.append(v)
                break
    personality = profile.get("character_personality", {}) or {}
    pd = personality.get("personality_details", {}) if isinstance(personality, dict) else {}
    if isinstance(pd, dict):
        traits = _text_list(pd.get("core_traits") or pd.get("positive_traits") or [])
        if traits:
            bits.append("Traits: " + ", ".join(traits[:5]))
    powers = profile.get("optional_powers_and_power_level") or profile.get("powers_and_abilities", {}) or {}
    if isinstance(powers, dict) and powers.get("is_enabled"):
        pdetails = powers.get("power_details", {}) if isinstance(powers.get("power_details"), dict) else {}
        pname = _safe(pdetails.get("power_name"))
        if pname:
            bits.append(f"Power: {pname}")
    return " - ".join(bits)


def _world_rules_lines(ms: dict) -> list[str]:
    """Pull selected world rules + their detail blocks from master_story."""
    block = ms.get("world_master_rules", {}) or {}
    if not isinstance(block, dict):
        return []
    selected = _text_list(block.get("selected", []))
    selected = [s for s in selected if s != "Custom"]
    rule_details = block.get("rule_details", {}) if isinstance(block.get("rule_details"), dict) else {}
    out: list[str] = []
    if selected:
        out.append(f"Selected: {', '.join(selected)}")
    detail_labels = [
        ("magic_rules", "Magic"),
        ("power_rules", "Powers"),
        ("demon_rules", "Demons"),
        ("monster_rules", "Monsters"),
        ("god_rules", "Gods"),
        ("technology_rules", "Technology"),
        ("race_species_rules", "Races / Species"),
        ("realm_dimension_rules", "Realms / Dimensions"),
        ("forbidden_rules", "Forbidden"),
        ("power_limits", "Power limits"),
    ]
    for key, label in detail_labels:
        text = _safe(rule_details.get(key))
        if text:
            out.append(f"{label}: {text}")
    return out


def _threats_lines(ms: dict) -> list[str]:
    """Full threat block: major, minor, sources, goals, stakes, time limit."""
    block = ms.get("major_threats_and_minor_side_threats", {}) or {}
    if not isinstance(block, dict):
        return []
    out: list[str] = []
    major = _safe(block.get("major_threat"))
    if major and major != "Custom":
        out.append(f"Major threat: {major}")
    minor = _text_list(block.get("minor_side_threats", []))
    minor = [m for m in minor if m != "Custom"]
    if minor:
        out.append(f"Minor threats: {', '.join(minor)}")
    details = block.get("threat_details", {}) if isinstance(block.get("threat_details"), dict) else {}
    for key, label in [
        ("main_threat_source", "Source"),
        ("main_threat_goal", "Goal"),
        ("main_threat_target", "Target"),
        ("stakes_if_major_threat_wins", "Stakes if it wins"),
        ("threat_level", "Level"),
        ("time_limit", "Time limit"),
        ("hidden_truth_behind_threat", "Hidden truth"),
    ]:
        v = _safe(details.get(key))
        if v:
            out.append(f"{label}: {v}")
    return out


def _relationships_lines(ch: dict) -> list[str]:
    rel_map = ch.get("character_relationship_map", {}) or {}
    if not isinstance(rel_map, dict):
        return []
    rels = _as_list(rel_map.get("relationships", []))
    out: list[str] = []
    for r in rels:
        if not isinstance(r, dict):
            continue
        pair = _safe(r.get("characters_involved"))
        rtype = _safe(r.get("relationship_change_type"))
        reason = _safe(r.get("reason"))
        if not pair:
            continue
        # Spec: character names lowercase
        pair = pair.lower()
        line = pair
        if rtype:
            line += f" - {rtype}"
        if reason:
            line += f" ({reason})"
        out.append(line)
    return out


def _build_id_to_name(ch: dict) -> dict[str, str]:
    """Map profile_id → character_name from characters.json for arc label resolution."""
    mapping: dict[str, str] = {}
    for p in _as_list(ch.get("created_major_character_profiles", [])) + _as_list(ch.get("created_side_character_profiles", [])):
        if isinstance(p, dict):
            pid = _safe(p.get("profile_id"))
            name = _safe(p.get("character_name"))
            if pid and name:
                mapping[pid] = name
    return mapping


def _plot_threads_lines(po: dict, ch: dict | None = None) -> list[str]:
    threads = po.get("plot_threads", {}) or {}
    if not isinstance(threads, dict):
        return []
    id_to_name: dict[str, str] = _build_id_to_name(ch) if ch else {}
    out: list[str] = []
    main = threads.get("main_plot_thread", {}) if isinstance(threads.get("main_plot_thread"), dict) else {}
    main_goal = _safe(main.get("goal"))
    if main_goal:
        out.append(f"Main goal: {main_goal}")
        for key, label in [("obstacles", "Obstacles"), ("turning_points", "Turning points"), ("resolution", "Resolution")]:
            v = main.get(key)
            text = ", ".join(_text_list(v)) if isinstance(v, list) else _safe(v)
            if text:
                out.append(f"  {label}: {text}")
    arcs = _as_list(threads.get("character_arc_threads", []))
    for arc in arcs:
        if not isinstance(arc, dict):
            continue
        cid = _safe(arc.get("character_id"))
        # Resolve ID → name; fall back to the raw value (may already be a name)
        label = id_to_name.get(cid, cid)
        # Spec: character names lowercase
        label = label.lower() if label else label
        start = _safe(arc.get("starting_state"))
        final = _safe(arc.get("final_state"))
        if label or start or final:
            out.append(f"Character arc {label}: {start or '?'} -> {final or '?'}")
    return out


def _assemble_story_lines(files: dict, all_scripts: list[dict] | None = None) -> list[str]:
    """Build a clean readable story document from the current official files."""
    lines: list[str] = []
    ms = files.get("master_story", {})
    ch = files.get("characters", {})
    po = files.get("plot_outline", {})
    cs = files.get("chapter_script", {})

    title = _story_title(ms)
    lines += [title, "=" * len(title), ""]

    # Subtitle: story types + foundation + ending direction
    story_types = _selected(ms.get("story_type"))
    foundation = _selected(ms.get("story_foundation"))
    ending = _selected(ms.get("ending_direction"))
    subtitle_bits = [b for b in [story_types, foundation, ending] if b]
    if subtitle_bits:
        lines += [" | ".join(subtitle_bits), ""]

    # SYNOPSIS — idea_so_far is the canonical premise field in the template
    premise = _safe(ms.get("idea_so_far"))
    if premise:
        lines += ["SYNOPSIS", "-" * 8, premise, ""]

    # WORLD — world_type + selected world rules + their detail blocks
    world_type = _selected(ms.get("world_type"))
    rules_lines = _world_rules_lines(ms)
    if world_type or rules_lines:
        lines += ["WORLD", "-" * 5]
        _append_field(lines, "Type", world_type)
        lines.extend(rules_lines)
        lines.append("")

    # FACTIONS — selected list
    factions_block = ms.get("major_factions_and_ruling_sides", {}) or {}
    factions = _text_list(factions_block.get("selected", []) if isinstance(factions_block, dict) else [])
    factions = [f for f in factions if f != "Custom"]
    if factions:
        lines += ["FACTIONS", "-" * 8, ", ".join(factions), ""]

    # FACTION VISUALS — visual signature per faction (injected into panel prompts)
    faction_vis_block = ms.get("faction_visual_signatures", {}) or {}
    faction_sigs = _as_list(faction_vis_block.get("signatures", []))
    faction_sigs = [s for s in faction_sigs if isinstance(s, dict) and _safe(s.get("faction_name"))]
    if faction_sigs:
        lines += ["FACTION VISUALS", "-" * 15]
        for sig in faction_sigs:
            name = _safe(sig.get("faction_name"))
            vis = _safe(sig.get("visual_signature"))
            pos = _safe(sig.get("positive_prompt"))
            neg = _safe(sig.get("negative_prompt"))
            if vis:
                lines.append(f"{name}: {vis}")
            if pos:
                lines.append(f"  AI prompt (positive): {pos}")
            if neg:
                lines.append(f"  AI prompt (negative): {neg}")
        lines.append("")

    # THREATS — full block (major + minor + source/goal/stakes/level)
    threat_lines = _threats_lines(ms)
    if threat_lines:
        lines += ["THREATS", "-" * 7]
        lines.extend(threat_lines)
        lines.append("")

    # CHARACTERS — real role + bio from canonical fields
    profiles = (ch.get("created_major_character_profiles", []) or []) + (ch.get("created_side_character_profiles", []) or [])
    if profiles:
        lines += ["CHARACTERS", "-" * 10]
        for profile in profiles:
            if not isinstance(profile, dict):
                continue
            name = _safe(profile.get("character_name") or profile.get("name"))
            if not name:
                continue
            # Spec: character names lowercase — parser silently drops fields otherwise
            name = name.lower()
            role = _profile_role(profile)
            bio = _profile_bio(profile)
            lines.append(f"{name}{' - ' + role if role else ''}")
            if bio:
                lines.append(f"  {bio}")
        lines.append("")

    # RELATIONSHIPS — character_relationship_map
    rel_lines = _relationships_lines(ch)
    if rel_lines:
        lines += ["RELATIONSHIPS", "-" * 13]
        lines.extend(rel_lines)
        lines.append("")

    # ARC OVERVIEW + PLOT THREADS
    story_arc = po.get("story_arc_overview", {}) or {}
    narrative_structure = _selected(po.get("narrative_structure"))
    arc_title = _safe(story_arc.get("arc_title"), "Current Arc")
    lines += ["ARC OVERVIEW", "-" * 12]
    _append_field(lines, "Structure", narrative_structure)
    _append_field(lines, "Arc", arc_title)
    _append_field(lines, "Arc length", _selected(story_arc.get("arc_length_type")))
    for label, key in [
        ("Summary", "arc_summary"),
        ("Story question", "main_story_question"),
        ("Emotional question", "central_emotional_question"),
        ("External conflict", "main_external_conflict"),
        ("Internal conflict", "main_internal_conflict"),
        ("Relationship conflict", "main_relationship_conflict"),
        ("Ending target", "ending_type_target"),
    ]:
        _append_field(lines, label, story_arc.get(key))
    lines.append("")

    thread_lines = _plot_threads_lines(po, ch)
    if thread_lines:
        lines += ["PLOT THREADS", "-" * 12]
        lines.extend(thread_lines)
        lines.append("")

    # CHAPTERS + per-chapter script content
    outline_chapters = _as_list((po.get("chapter_or_episode_list", {}) or {}).get("chapters", []))
    scene_cards = _as_list((po.get("scene_cards", {}) or {}).get("scenes", []))
    scenes_by_chapter = _group_scenes_by_chapter(scene_cards)

    # Build chapter_id → pages map. Prefer all_scripts (cross-version aggregate) so every
    # approved chapter's panel dialogue/narration appears in the document, not just the one
    # chapter currently loaded in the working slot.
    if all_scripts:
        script_pages_by_chapter: dict[str, list[dict]] = {
            _safe((s.get("chapter_metadata") or {}).get("chapter_id")): _as_list(s.get("pages"))
            for s in all_scripts
            if _safe((s.get("chapter_metadata") or {}).get("chapter_id"))
        }
    else:
        # Fallback: single working-slot chapter
        cid = _safe((cs.get("chapter_metadata") or {}).get("chapter_id"))
        script_pages_by_chapter = {cid: _as_list(cs.get("pages"))} if cid else {}

    if not outline_chapters:
        lines += ["CHAPTERS", "=" * 8, "", "No chapters found. Create chapters on the Plot Board first.", ""]
        return lines

    lines += ["FULL STORY", "=" * 10, ""]
    for chapter in sorted(outline_chapters, key=lambda c: int(c.get("chapter_number") or 0) if isinstance(c, dict) else 0):
        if not isinstance(chapter, dict):
            continue
        header = _chapter_header(chapter)
        chapter_id = _safe(chapter.get("chapter_id"))
        chapter_scenes = scenes_by_chapter.get(chapter_id, [])
        chapter_script_pages = script_pages_by_chapter.get(chapter_id, [])

        lines += [header, "-" * len(header)]
        _append_field(lines, "Arc", chapter.get("arc_title") or arc_title)
        _append_field(lines, "Structure beat", chapter.get("structure_section"))
        lines.append("")

        lines += ["Story", "-----"]
        script_lines = _panel_story_lines(chapter_script_pages)
        if script_lines:
            lines.extend(script_lines)
        else:
            for key in [
                "chapter_purpose",
                "summary",
                "main_conflict",
                "emotional_beat",
                "twist_or_hook",
                "ending_cliffhanger",
                "custom_chapter_details",
            ]:
                text = _safe(chapter.get(key))
                if text:
                    lines.append(text)
            if chapter_scenes:
                lines.append("")
                lines.append("Key scenes")
                seen_scene_beats: set[str] = set()
                for scene in chapter_scenes:
                    scene_order = _safe(scene.get("scene_order") or scene.get("scene_id"), "?")
                    scene_goal = _safe(scene.get("scene_goal"))
                    scene_reveal = _safe(scene.get("new_information_revealed"))
                    scene_ending = _safe(scene.get("ending_beat"))
                    scene_bits = [scene_goal]
                    if scene_reveal and scene_reveal != scene_goal:
                        scene_bits.append(scene_reveal)
                    if scene_ending and scene_ending not in scene_bits:
                        scene_bits.append(scene_ending)
                    scene_text = " ".join(bit for bit in scene_bits if bit)
                    if scene_text in seen_scene_beats:
                        continue
                    seen_scene_beats.add(scene_text)
                    if scene_text:
                        lines.append(f"{scene_order}. {scene_text}")
        lines.append("")

        lines += ["Chapter Review", "--------------"]
        _append_field(lines, "Characters active", chapter.get("characters_present"))
        _append_field(lines, "Relationship movement", chapter.get("relationships_used"))
        _append_field(lines, "Threat / faction movement", _text_list(chapter.get("threats_used")) + _text_list(chapter.get("factions_used")))
        _append_field(lines, "World or powers", _text_list(chapter.get("world_rules_shown")) + _text_list(chapter.get("power_system_shown")))
        _append_field(lines, "Open hook", chapter.get("ending_cliffhanger") or chapter.get("twist_or_hook"))
        lines.append("")

    return lines


def _script_pages_by_scene(cs: dict) -> dict[str, list[dict]]:
    """Group chapter_script pages by scene_id."""
    grouped: dict[str, list[dict]] = {}
    for page in _as_list(cs.get("pages", [])):
        if not isinstance(page, dict):
            continue
        sid = _safe(page.get("scene_id"))
        if sid:
            grouped.setdefault(sid, []).append(page)
    return grouped


def _scene_script_lines(pages: list[dict], indent: str = "  ") -> list[str]:
    """Render the script content for one scene: per-panel narration, dialogue, SFX, visuals."""
    out: list[str] = []
    for page in pages:
        page_label = _safe(page.get("page_number") or page.get("page_id"))
        out.append(f"{indent}- Page {page_label}")
        for panel in _as_list(page.get("panels", [])):
            if not isinstance(panel, dict):
                continue
            panel_no = _safe(panel.get("panel_number") or panel.get("panel_id"))
            shot_block = panel.get("camera_shot")
            shot = _selected(shot_block) if isinstance(shot_block, dict) else _safe(shot_block)
            visual = _safe(panel.get("visual"))
            action = _safe(panel.get("character_action"))
            narration = _safe(panel.get("narration"))
            header = f"{indent}  Panel {panel_no}"
            if shot:
                header += f" [{shot}]"
            out.append(header)
            if visual:
                out.append(f"{indent}    Visual: {visual}")
            if action:
                out.append(f"{indent}    Action: {action}")
            if narration:
                out.append(f"{indent}    Narration: {narration}")
            for d in _as_list(panel.get("dialogue", [])):
                if not isinstance(d, dict):
                    continue
                text = _safe(d.get("text"))
                if not text:
                    continue
                speaker = _safe(d.get("speaker_name") or d.get("speaker") or d.get("speaker_id")) or "?"
                bubble_block = d.get("speech_bubble_type")
                bubble = _selected(bubble_block) if isinstance(bubble_block, dict) else _safe(bubble_block)
                # Skip Narrator entries — Narration: field already covers caption text
                if speaker.lower() in ("narrator", "narration") or (bubble and bubble.lower() == "narration"):
                    continue
                bubble_tag = f" ({bubble})" if bubble and bubble != "Normal" else ""
                out.append(f'{indent}    {speaker}{bubble_tag}: "{text}"')
            sfx_items = _as_list(panel.get("sound_effects", []))
            sfx_texts = []
            for s in sfx_items:
                if isinstance(s, dict):
                    t = _safe(s.get("sfx_text"))
                    if t:
                        sfx_texts.append(t)
                elif isinstance(s, str) and s.strip():
                    sfx_texts.append(s.strip())
            if sfx_texts:
                out.append(f"{indent}    SFX: {', '.join(sfx_texts)}")
    return out


def _build_pages_by_scene_all(all_scripts: list[dict]) -> dict[str, list[dict]]:
    """Aggregate pages_by_scene across ALL chapter scripts (latest chapter wins for same scene_id)."""
    merged: dict[str, list[dict]] = {}
    for cs in all_scripts:
        for sid, pages in _script_pages_by_scene(cs).items():
            merged[sid] = pages  # later scripts overwrite earlier; all_scripts sorted by chapter_number
    return merged


def _assemble_scenes_lines(files: dict, all_scripts: list[dict] | None = None) -> list[str]:
    """Build full scene-card export with chapter-script dialogue/SFX/visuals injected per scene."""
    po = files.get("plot_outline", {})
    ms = files.get("master_story", {}) or {}
    cs = files.get("chapter_script", {}) or {}
    lines: list[str] = []

    title = _story_title(ms, "Story")
    lines += [f"SCENES - {title}", "=" * (9 + len(title)), ""]

    chapters = _as_list((po.get("chapter_or_episode_list", {}) or {}).get("chapters", []))
    scene_cards = _as_list((po.get("scene_cards", {}) or {}).get("scenes", []))
    scenes_by_chapter = _group_scenes_by_chapter(scene_cards)

    # Use all chapter scripts if provided; otherwise fall back to the single current one
    if all_scripts:
        pages_by_scene = _build_pages_by_scene_all(all_scripts)
        chapters_with_script: set[str] = {
            _safe((s.get("chapter_metadata") or {}).get("chapter_id"))
            for s in all_scripts
        }
    else:
        pages_by_scene = _script_pages_by_scene(cs)
        chapters_with_script = {_safe((cs.get("chapter_metadata") or {}).get("chapter_id"))}

    if not chapters and not scene_cards:
        lines.append("No scenes found. Add scenes in the Scene Cards studio.")
        return lines

    chapter_ids_seen: set[str] = set()
    for chapter in sorted(chapters, key=lambda c: int(c.get("chapter_number") or 0) if isinstance(c, dict) else 0):
        if not isinstance(chapter, dict):
            continue
        chapter_id = _safe(chapter.get("chapter_id"))
        chapter_ids_seen.add(chapter_id)
        header = _chapter_header(chapter)
        lines += [header, "-" * len(header)]
        _append_field(lines, "Chapter purpose", chapter.get("chapter_purpose"))
        _append_field(lines, "Chapter summary", chapter.get("summary"))
        _append_field(lines, "Chapter conflict", chapter.get("main_conflict"))
        _append_field(lines, "Chapter hook", chapter.get("ending_cliffhanger") or chapter.get("twist_or_hook"))
        if chapter_id and chapter_id in chapters_with_script:
            lines.append("Script: AVAILABLE (panel content shown per scene below)")
        elif chapter_id:
            lines.append("Script: not generated for this chapter yet")

        chapter_scenes = scenes_by_chapter.get(chapter_id, [])
        if not chapter_scenes:
            lines.append("No scene cards for this chapter.")
            lines.append("")
            continue

        lines.append("")
        for scene in chapter_scenes:
            scene_no = _safe(scene.get("scene_order") or scene.get("scene_id"), "?")
            scene_title = _safe(scene.get("scene_title") or scene.get("scene_id"), f"Scene {scene_no}")
            scene_id = _safe(scene.get("scene_id"))
            lines.append(f"Scene {scene_no}: {scene_title}")
            _append_field(lines, "Location", scene.get("location"), "  ")
            _append_field(lines, "Time", scene.get("time"), "  ")
            _append_field(lines, "Characters", scene.get("characters_present"), "  ")
            _append_field(lines, "Goal", scene.get("scene_goal"), "  ")
            _append_field(lines, "Conflict", scene.get("scene_conflict"), "  ")
            _append_field(lines, "Relationship dynamic", scene.get("relationship_dynamic_used"), "  ")
            _append_field(lines, "New information", scene.get("new_information_revealed"), "  ")
            _append_field(lines, "Action/dialogue focus", scene.get("action_or_dialogue_focus"), "  ")
            _append_field(lines, "Visual manga moment", scene.get("visual_manga_moment"), "  ")
            _append_field(lines, "Panel mood", scene.get("panel_mood"), "  ")
            _append_field(lines, "Ending beat", scene.get("ending_beat"), "  ")
            _append_field(lines, "Custom details", scene.get("custom_scene_details"), "  ")

            # Inject script content if this scene has generated pages.
            scene_pages = pages_by_scene.get(scene_id, []) if scene_id else []
            if scene_pages:
                lines.append("  Script:")
                lines.extend(_scene_script_lines(scene_pages, indent="  "))
            lines.append("")

    orphan_scenes = [scene for scene in scene_cards if isinstance(scene, dict) and _safe(scene.get("chapter_id")) not in chapter_ids_seen]
    if orphan_scenes:
        lines += ["Unassigned Scenes", "-" * 17]
        for scene in orphan_scenes:
            lines.append(_safe(scene.get("scene_id"), "Scene"))
            _append_field(lines, "Chapter ID", scene.get("chapter_id"), "  ")
            _append_field(lines, "Goal", scene.get("scene_goal"), "  ")
            _append_field(lines, "Visual manga moment", scene.get("visual_manga_moment"), "  ")
            lines.append("")

    return lines


# ─── validation ─────────────────────────────────────────────────────────────

def _validate_export(files: dict, all_scripts: list[dict]) -> dict:
    """Surface data-quality issues that the export tool can't fix on its own.

    Each warning has level (critical/high/medium/info), category, message, where.
    Designed to mirror the G-Ink Studio compatibility requirements so the user
    sees what's missing before downloading.
    """
    warnings: list[dict] = []
    ms = files.get("master_story", {}) or {}
    ch = files.get("characters", {}) or {}
    po = files.get("plot_outline", {}) or {}

    # 1. Pages-per-chapter sanity (missing scene_cards causes single-page chapters)
    chapters = _as_list((po.get("chapter_or_episode_list", {}) or {}).get("chapters", []))
    scene_cards = _as_list((po.get("scene_cards", {}) or {}).get("scenes", []))
    scenes_by_chapter: dict[str, list] = {}
    for sc in scene_cards:
        if isinstance(sc, dict):
            cid = _safe(sc.get("chapter_id"))
            if cid:
                scenes_by_chapter.setdefault(cid, []).append(sc)

    pages_by_chapter: dict[str, list] = {}
    for cs in all_scripts:
        cid = _safe((cs.get("chapter_metadata") or {}).get("chapter_id"))
        if cid:
            pages_by_chapter[cid] = _as_list(cs.get("pages"))

    for chap in chapters:
        if not isinstance(chap, dict):
            continue
        cid = _safe(chap.get("chapter_id"))
        ch_num = _safe(chap.get("chapter_number")) or "?"
        scene_count = len(scenes_by_chapter.get(cid, []))
        page_count = len(pages_by_chapter.get(cid, []))
        if scene_count >= 2 and page_count <= 1:
            warnings.append({
                "level": "high",
                "category": "missing_pages",
                "message": (
                    f"Chapter {ch_num} has {scene_count} scenes defined but only {page_count} page(s) in script. "
                    "Regenerate the script after scenes are populated to expand to one page per scene."
                ),
                "where": "Studio → Manga Script → Generate (per chapter)",
            })

    # 2. Speakers without character profiles
    profiles = _as_list(ch.get("created_major_character_profiles", [])) + _as_list(ch.get("created_side_character_profiles", []))
    profile_names_lower: set[str] = set()
    for p in profiles:
        if isinstance(p, dict):
            n = _safe(p.get("character_name") or p.get("name"))
            if n:
                profile_names_lower.add(n.lower())

    speakers_seen: set[str] = set()
    for cs in all_scripts:
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            for panel in _as_list(page.get("panels")):
                if not isinstance(panel, dict):
                    continue
                for d in _as_list(panel.get("dialogue", [])):
                    if not isinstance(d, dict):
                        continue
                    speaker = _safe(d.get("speaker_name") or d.get("speaker") or d.get("speaker_id"))
                    if speaker:
                        speakers_seen.add(speaker.lower())

    skip_speakers = {"narrator", "narration", "?", ""}
    missing_speakers = sorted(s for s in speakers_seen if s not in skip_speakers and s not in profile_names_lower)
    for speaker in missing_speakers:
        warnings.append({
            "level": "high",
            "category": "missing_profile",
            "message": (
                f"Speaker '{speaker}' has dialogue but no character profile. "
                "G-Ink Studio will skip this character; add a major or side profile."
            ),
            "where": "Studio → Cast (major) or Side Characters",
        })

    # 3. FACTION VISUALS missing when factions defined
    factions_block = ms.get("major_factions_and_ruling_sides", {}) or {}
    factions = _text_list(factions_block.get("selected", []) if isinstance(factions_block, dict) else [])
    factions = [f for f in factions if f and f != "Custom"]
    sigs = _as_list((ms.get("faction_visual_signatures", {}) or {}).get("signatures", []))
    sigs = [s for s in sigs if isinstance(s, dict) and _safe(s.get("faction_name")) and _safe(s.get("visual_signature"))]
    if factions and not sigs:
        warnings.append({
            "level": "medium",
            "category": "missing_faction_visuals",
            "message": (
                f"{len(factions)} faction(s) defined but no FACTION VISUALS populated. "
                "Without visual signatures, faction-specific gear can't be injected into panel prompts."
            ),
            "where": "Studio → Faction Visuals",
        })

    # 4. LOCATIONS section empty
    locs = _as_list((po.get("locations", {}) or {}).get("locations", []))
    if not locs:
        warnings.append({
            "level": "critical",
            "category": "missing_locations",
            "message": (
                "No LOCATIONS defined. The studio falls back to keyword heuristics; "
                "AI prompts will be inconsistent across panels."
            ),
            "where": "Studio → Locations",
        })

    # 5. Pages without location_id (≥30% threshold)
    pages_without_loc = 0
    total_pages = 0
    for cs in all_scripts:
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            total_pages += 1
            if not _safe(page.get("location_id")):
                pages_without_loc += 1
    if total_pages > 0 and pages_without_loc / total_pages >= 0.3:
        warnings.append({
            "level": "medium",
            "category": "missing_location_id",
            "message": (
                f"{pages_without_loc}/{total_pages} pages have no location_id. "
                "Run 'Generate All' in Visuals Studio so each panel resolves to a named location."
            ),
            "where": "Studio → Visuals Studio",
        })

    # 5b. Pages whose declared Location isn't mentioned in any of their panels
    #     (BUNDLE-AUDIT #1 — the scene→location map looks scrambled). The export
    #     repairs the header when it can; this flags it so the source gets fixed.
    loc_by_id = _build_loc_by_id(po)
    all_loc_names = {
        _safe(l.get("name")) for l in _as_list((po.get("locations") or {}).get("locations", []))
    } | {_safe(v.get("name")) for v in loc_by_id.values()}
    all_loc_names = {n for n in all_loc_names if len(n) >= 4}
    mismatched: list[str] = []
    for cs in all_scripts:
        ch_num = _safe((cs.get("chapter_metadata") or {}).get("chapter_number") or (cs.get("chapter_metadata") or {}).get("chapter_id"))
        scene_loc_name = {
            _safe(s.get("scene_id")): _safe(s.get("location"))
            for s in _as_list(cs.get("chapter_scene_breakdown")) if isinstance(s, dict) and s.get("location")
        }
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            lid = _safe(page.get("location_id"))
            name = _safe(loc_by_id.get(lid, {}).get("name")) if lid in loc_by_id else scene_loc_name.get(_safe(page.get("scene_id")), "")
            if len(name) < 4:
                continue
            ptext = " ".join(
                f"{_safe(pn.get('visual'))} {_safe(pn.get('background_details'))} {_safe(pn.get('character_action'))}"
                for pn in _as_list(page.get("panels")) if isinstance(pn, dict)
            ).lower()
            if ptext and name.lower() not in ptext and any(n.lower() in ptext for n in all_loc_names if n != name):
                mismatched.append(f"Ch{ch_num} p{_safe(page.get('page_number'))} (says {name!r})")
    if mismatched:
        warnings.append({
            "level": "high",
            "category": "location_mismatch",
            "message": (
                "Page Location(s) don't match the panel content — the scene→location map looks scrambled: "
                + "; ".join(mismatched[:8]) + (f" (+{len(mismatched) - 8} more)" if len(mismatched) > 8 else "")
                + ". The export auto-corrects the header where it can; fix the scene→location assignment at the source so Purpose / Location / panel visuals agree."
            ),
            "where": "Studio → Plot Board (scene cards) / Visuals Studio",
        })

    # 6. Character names not lowercased in source data (info — export auto-lowers)
    upper_names: list[str] = []
    for p in profiles:
        if isinstance(p, dict):
            n = _safe(p.get("character_name") or p.get("name"))
            if n and n != n.lower():
                upper_names.append(n)
    if upper_names:
        sample = ", ".join(upper_names[:5]) + (f" (+{len(upper_names) - 5} more)" if len(upper_names) > 5 else "")
        warnings.append({
            "level": "info",
            "category": "name_case",
            "message": (
                f"Character profile name(s) not lowercased: {sample}. "
                "Export auto-lowercases for studio compatibility, but consider renaming in the source."
            ),
            "where": "Studio → Cast",
        })

    # 7. Visual prompts that carry colour / lighting / style noise (info —
    #    export now auto-cleans these and prepends "black and white Japanese
    #    manga style", but it's worth flagging so the source gets fixed).
    def _is_dirty(text: object) -> bool:
        raw = " ".join(str(text or "").split())
        return bool(raw) and sanitize_visual_prompt(raw) != raw and sanitize_visual_prompt(raw) != ""
    dirty_chars: list[str] = []
    for p in profiles:
        if not isinstance(p, dict):
            continue
        det = _appearance_block(p)
        if _is_dirty(det.get("ai_image_prompt_notes")):
            dirty_chars.append(_safe(p.get("character_name") or p.get("name")))
    dirty_locs: list[str] = []
    for loc in _as_list((po.get("locations") or {}).get("locations", [])):
        if isinstance(loc, dict) and _is_dirty(loc.get("positive_prompt")):
            dirty_locs.append(_safe(loc.get("name")))
    if dirty_chars or dirty_locs:
        bits = []
        if dirty_chars:
            bits.append("character sheet(s): " + ", ".join(filter(None, dirty_chars[:6])))
        if dirty_locs:
            bits.append("location(s): " + ", ".join(filter(None, dirty_locs[:6])))
        warnings.append({
            "level": "info",
            "category": "prompt_noise",
            "message": (
                "AI image prompt(s) contain colour / lighting / 'cinematic'-style words or a per-entity "
                "style tag — " + "; ".join(bits) + ". The export strips these and prepends "
                "'black and white Japanese manga style' automatically; consider cleaning the source so the "
                "stored prompt is already a short visual-only descriptor list."
            ),
            "where": "Studio → Cast / Locations / Faction Visuals",
        })

    return {"warnings": warnings, "count": len(warnings)}


def _format_validation_lines(report: dict) -> list[str]:
    """Render the validation report as a list of markdown-friendly lines."""
    lines: list[str] = ["VALIDATION REPORT", "=" * 17, ""]
    warnings = _as_list(report.get("warnings"))
    if not warnings:
        lines.append("No issues detected. Story export is fully compatible with G-Ink Studio.")
        return lines

    lines.append(f"{len(warnings)} issue(s) detected. Review below before importing into G-Ink Studio.")
    lines.append("")

    by_level: dict[str, list[dict]] = {}
    for w in warnings:
        by_level.setdefault(_safe(w.get("level"), "info"), []).append(w)

    for level in ["critical", "high", "medium", "info"]:
        items = by_level.get(level, [])
        if not items:
            continue
        lines.append(f"{level.upper()} ({len(items)})")
        lines.append("-" * (len(level) + len(str(len(items))) + 3))
        for w in items:
            cat = _safe(w.get("category"))
            msg = _safe(w.get("message"))
            where = _safe(w.get("where"))
            lines.append(f"[{cat}] {msg}")
            if where:
                lines.append(f"  Where to fix: {where}")
            lines.append("")
    return lines
