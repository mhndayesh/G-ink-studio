from __future__ import annotations

"""Low-level shared helpers for the export renderers.

Safe scalar/list coercion, story/chapter title helpers, the lines->text/markdown/docx
formatters, and the file-loading helpers (_get_all_files, _all_chapter_scripts) plus a
few small lookups (_appearance_block, _build_loc_by_id, _build_id_to_name). Pure logic,
no FastAPI types. Shared by export_service / export_validation / export_character_files.
"""

import io
import json
import zipfile

from app.repositories.sqlite_registry import SQLiteRegistry


def story_safe_title(files: dict, fallback: str = "story") -> tuple[str, str]:
    """Return (display_title, filesystem_safe_title) from the master_story file."""
    ms = files.get("master_story", {}) or {}
    title = _safe(ms.get("title") or ms.get("story_title"), fallback)
    safe = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip().replace(" ", "_")
    return title, safe


# ─── helpers ────────────────────────────────────────────────────────────────


def _get_all_files(story_id: str, registry: SQLiteRegistry) -> dict:
    out: dict = {}
    for ft in ["master_story", "characters", "plot_outline", "plot_workspace", "chapter_script"]:
        rec = registry.get_current_file(story_id, ft)
        if rec:
            out[ft] = rec.get("json_copy", {})
    return out


def _safe(v: object, fallback: str = "") -> str:
    if v is None:
        return fallback
    return str(v).strip() or fallback


def _selected(v: object) -> str:
    if isinstance(v, dict):
        selected = v.get("selected")
        if isinstance(selected, list):
            return ", ".join(_text_list(selected))
        return _safe(selected)
    if isinstance(v, list):
        return ", ".join(_text_list(v))
    return _safe(v)


def _as_list(v: object) -> list:
    return v if isinstance(v, list) else []


def _text_list(v: object) -> list[str]:
    out: list[str] = []
    for item in _as_list(v):
        if isinstance(item, str):
            text = _safe(item)
        elif isinstance(item, dict):
            text = _safe(
                item.get("name")
                or item.get("character_name")
                or item.get("profile_name")
                or item.get("title")
                or item.get("label")
                or item.get("summary")
                or item.get("description")
                or item.get("relationship_id")
                or item.get("threat_id_or_name")
            )
        else:
            text = _safe(item)
        if text:
            out.append(text)
    return out


def _append_field(lines: list[str], label: str, value: object, indent: str = "") -> None:
    if isinstance(value, list):
        text = ", ".join(_text_list(value))
    else:
        text = _safe(value)
    if text:
        lines.append(f"{indent}{label}: {text}")


def _story_title(ms: dict, fallback: str = "Untitled Story") -> str:
    foundation = ms.get("story_foundation") if isinstance(ms.get("story_foundation"), dict) else {}
    return _safe(ms.get("title") or ms.get("story_title") or foundation.get("story_title"), fallback)


def _chapter_header(chapter: dict) -> str:
    ch_num = _safe(chapter.get("chapter_number") or chapter.get("chapter_id"), "?")
    ch_title = _safe(chapter.get("chapter_title") or chapter.get("title"))
    return f"Chapter {ch_num}{': ' + ch_title if ch_title else ''}"


def _lines_to_text(lines: list[str]) -> str:
    return "\n".join(lines) + "\n"


def _lines_to_markdown(lines: list[str]) -> str:
    """Convert the assembled lines to Markdown with proper heading levels."""
    md: list[str] = []
    for i, line in enumerate(lines):
        # Detect heading markers based on next-line underlines
        if i + 1 < len(lines):
            nxt = lines[i + 1]
            if nxt and all(c == "=" for c in nxt) and len(nxt) >= 3:
                md.append(f"# {line}")
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= 3:
                md.append(f"## {line}")
                continue
        # Skip the underline rows themselves
        if line and (all(c == "=" for c in line) or all(c == "-" for c in line)):
            md.append("")
            continue
        md.append(line)
    return "\n".join(md) + "\n"


def _lines_to_docx(lines: list[str], title: str) -> bytes:
    """Generate a .docx file from assembled lines. Falls back to plain bytes on import error."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[attr-defined]

        doc = Document()

        # Remove default empty paragraph
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        i = 0
        while i < len(lines):
            line = lines[i]
            nxt = lines[i + 1] if i + 1 < len(lines) else ""

            if nxt and all(c == "=" for c in nxt) and len(nxt) >= 3:
                doc.add_heading(line, level=1)
                i += 2  # skip underline
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= 3:
                doc.add_heading(line, level=2)
                i += 2
                continue
            if line and (all(c == "=" for c in line) or all(c == "-" for c in line)):
                i += 1
                continue
            if line == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # python-docx not installed — return plain text bytes with .docx extension
        return _lines_to_text(lines).encode("utf-8")


# ─── visuals assembly ────────────────────────────────────────────────────────


def _all_chapter_scripts(story_id: str, registry: SQLiteRegistry) -> list[dict]:
    """Walk every version's chapter_script.json snapshot, deduped by chapter_id.

    Latest version per chapter wins (so an updated chapter overrides older snapshots).
    """
    rows = registry.list_files_across_versions(story_id, "chapter_script")
    by_chapter: dict[str, dict] = {}
    for row in rows:
        data = row.get("json_copy", {}) or {}
        if not _as_list(data.get("pages")):
            continue
        ch_id = _safe((data.get("chapter_metadata") or {}).get("chapter_id")) or _safe(row.get("version_id"))
        if not ch_id:
            continue
        # Latest wins (rows are ordered by version_number ASC).
        by_chapter[ch_id] = data
    # Sort by chapter_number
    return sorted(
        by_chapter.values(),
        key=lambda d: int((d.get("chapter_metadata") or {}).get("chapter_number") or 0),
    )


def _appearance_block(profile: dict) -> dict:
    block = profile.get("appearance_and_visual_design", {})
    if isinstance(block, dict):
        details = block.get("appearance_details", {})
        if isinstance(details, dict) and details:
            return details
    # Fallback: some profiles may flatten the fields onto the root.
    return profile if any(profile.get(k) for k in ("hair_style", "main_outfit_description", "color_palette")) else {}


def _build_loc_by_id(po: dict) -> dict[str, dict]:
    """Map location_id → location dict from plot_outline."""
    locs_block = (po.get("locations") or {})
    out: dict[str, dict] = {}
    for loc in _as_list(locs_block.get("locations", [])):
        if isinstance(loc, dict):
            lid = _safe(loc.get("location_id"))
            if lid:
                out[lid] = loc
    return out


def _build_id_to_name(ch: dict) -> dict[str, str]:
    """Map profile_id → character_name from characters.json for arc label resolution."""
    mapping: dict[str, str] = {}
    for p in _as_list(ch.get("created_major_character_profiles", [])) + _as_list(ch.get("created_side_character_profiles", [])):
        if isinstance(p, dict):
            pid = _safe(p.get("profile_id"))
            name = _safe(p.get("character_name"))
            if pid and name:
                mapping[pid] = name
    return mapping
