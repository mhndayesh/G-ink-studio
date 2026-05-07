from __future__ import annotations

import io
import json
import zipfile
from typing import Literal

from fastapi import APIRouter, Depends, Query
from fastapi.responses import Response, StreamingResponse

from app.core.auth import require_story_access
from app.main_dependencies import get_registry
from app.repositories.sqlite_registry import SQLiteRegistry

router = APIRouter(
    dependencies=[Depends(require_story_access)],
    prefix="/stories/{story_id}/export",
    tags=["export"],
)

# ─── helpers ────────────────────────────────────────────────────────────────

def _get_all_files(story_id: str, registry: SQLiteRegistry) -> dict:
    out: dict = {}
    for ft in ["master_story", "characters", "plot_outline", "plot_workspace", "chapter_script"]:
        rec = registry.get_current_file(story_id, ft)
        if rec:
            out[ft] = rec.get("json_copy", {})
    return out


def _safe(v: object, fallback: str = "") -> str:
    if v is None:
        return fallback
    return str(v).strip() or fallback


def _selected(v: object) -> str:
    if isinstance(v, dict):
        selected = v.get("selected")
        if isinstance(selected, list):
            return ", ".join(_text_list(selected))
        return _safe(selected)
    if isinstance(v, list):
        return ", ".join(_text_list(v))
    return _safe(v)


def _as_list(v: object) -> list:
    return v if isinstance(v, list) else []


def _text_list(v: object) -> list[str]:
    out: list[str] = []
    for item in _as_list(v):
        if isinstance(item, str):
            text = _safe(item)
        elif isinstance(item, dict):
            text = _safe(
                item.get("name")
                or item.get("character_name")
                or item.get("profile_name")
                or item.get("title")
                or item.get("label")
                or item.get("summary")
                or item.get("description")
                or item.get("relationship_id")
                or item.get("threat_id_or_name")
            )
        else:
            text = _safe(item)
        if text:
            out.append(text)
    return out


def _append_field(lines: list[str], label: str, value: object, indent: str = "") -> None:
    if isinstance(value, list):
        text = ", ".join(_text_list(value))
    else:
        text = _safe(value)
    if text:
        lines.append(f"{indent}{label}: {text}")


def _story_title(ms: dict, fallback: str = "Untitled Story") -> str:
    foundation = ms.get("story_foundation") if isinstance(ms.get("story_foundation"), dict) else {}
    return _safe(ms.get("title") or ms.get("story_title") or foundation.get("story_title"), fallback)


def _chapter_header(chapter: dict) -> str:
    ch_num = _safe(chapter.get("chapter_number") or chapter.get("chapter_id"), "?")
    ch_title = _safe(chapter.get("chapter_title") or chapter.get("title"))
    return f"Chapter {ch_num}{': ' + ch_title if ch_title else ''}"


def _group_scenes_by_chapter(scene_cards: list[dict]) -> dict[str, list[dict]]:
    grouped: dict[str, list[dict]] = {}
    for scene in scene_cards:
        if not isinstance(scene, dict):
            continue
        grouped.setdefault(_safe(scene.get("chapter_id")), []).append(scene)
    for scenes in grouped.values():
        scenes.sort(key=lambda s: int(s.get("scene_order") or 0))
    return grouped


def _panel_story_lines(pages: list[dict]) -> list[str]:
    lines: list[str] = []
    for page in pages:
        if not isinstance(page, dict):
            continue
        page_label = _safe(page.get("page_number") or page.get("page_id"))
        page_bits = []
        if _safe(page.get("page_purpose")):
            page_bits.append(_safe(page.get("page_purpose")))
        if _safe(page.get("page_mood")):
            page_bits.append(f"Mood: {_safe(page.get('page_mood'))}")
        if page_bits:
            lines.append(f"Page {page_label}: " + " ".join(page_bits))
        for panel in _as_list(page.get("panels")):
            if not isinstance(panel, dict):
                continue
            bits = [
                _safe(panel.get("visual")),
                _safe(panel.get("character_action")),
                _safe(panel.get("narration")),
            ]
            for dialogue in _as_list(panel.get("dialogue")):
                if not isinstance(dialogue, dict):
                    continue
                text = _safe(dialogue.get("text"))
                if not text:
                    continue
                speaker = _safe(dialogue.get("speaker_name") or dialogue.get("speaker") or dialogue.get("speaker_id"))
                bits.append(f"{speaker}: {text}" if speaker else text)
            panel_text = " ".join(bit for bit in bits if bit)
            if panel_text:
                panel_no = _safe(panel.get("panel_number") or panel.get("panel_id"))
                lines.append(f"Panel {panel_no}: {panel_text}")
    return lines


# ─── story assembly ──────────────────────────────────────────────────────────

def _lines_to_text(lines: list[str]) -> str:
    return "\n".join(lines) + "\n"


def _lines_to_markdown(lines: list[str]) -> str:
    """Convert the assembled lines to Markdown with proper heading levels."""
    md: list[str] = []
    for i, line in enumerate(lines):
        # Detect heading markers based on next-line underlines
        if i + 1 < len(lines):
            nxt = lines[i + 1]
            if nxt and all(c == "=" for c in nxt) and len(nxt) >= 3:
                md.append(f"# {line}")
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= 3:
                md.append(f"## {line}")
                continue
        # Skip the underline rows themselves
        if line and (all(c == "=" for c in line) or all(c == "-" for c in line)):
            md.append("")
            continue
        md.append(line)
    return "\n".join(md) + "\n"


def _lines_to_docx(lines: list[str], title: str) -> bytes:
    """Generate a .docx file from assembled lines. Falls back to plain bytes on import error."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[attr-defined]

        doc = Document()

        # Remove default empty paragraph
        for p in doc.paragraphs:
            p._element.getparent().remove(p._element)

        i = 0
        while i < len(lines):
            line = lines[i]
            nxt = lines[i + 1] if i + 1 < len(lines) else ""

            if nxt and all(c == "=" for c in nxt) and len(nxt) >= 3:
                doc.add_heading(line, level=1)
                i += 2  # skip underline
                continue
            if nxt and all(c == "-" for c in nxt) and len(nxt) >= 3:
                doc.add_heading(line, level=2)
                i += 2
                continue
            if line and (all(c == "=" for c in line) or all(c == "-" for c in line)):
                i += 1
                continue
            if line == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        # python-docx not installed — return plain text bytes with .docx extension
        return _lines_to_text(lines).encode("utf-8")


# ─── visuals assembly ────────────────────────────────────────────────────────

def _all_chapter_scripts(story_id: str, registry: SQLiteRegistry) -> list[dict]:
    """Walk every version's chapter_script.json snapshot, deduped by chapter_id.

    Latest version per chapter wins (so an updated chapter overrides older snapshots).
    """
    rows = registry.list_files_across_versions(story_id, "chapter_script")
    by_chapter: dict[str, dict] = {}
    for row in rows:
        data = row.get("json_copy", {}) or {}
        if not _as_list(data.get("pages")):
            continue
        ch_id = _safe((data.get("chapter_metadata") or {}).get("chapter_id")) or _safe(row.get("version_id"))
        if not ch_id:
            continue
        # Latest wins (rows are ordered by version_number ASC).
        by_chapter[ch_id] = data
    # Sort by chapter_number
    return sorted(
        by_chapter.values(),
        key=lambda d: int((d.get("chapter_metadata") or {}).get("chapter_number") or 0),
    )


def _appearance_block(profile: dict) -> dict:
    block = profile.get("appearance_and_visual_design", {})
    if isinstance(block, dict):
        details = block.get("appearance_details", {})
        if isinstance(details, dict) and details:
            return details
    # Fallback: some profiles may flatten the fields onto the root.
    return profile if any(profile.get(k) for k in ("hair_style", "main_outfit_description", "color_palette")) else {}


def _character_reference_lines(characters: dict) -> list[str]:
    """Per-character visual reference sheet for the artist."""
    lines: list[str] = []
    profiles = (characters.get("created_major_character_profiles", []) or []) + (characters.get("created_side_character_profiles", []) or [])
    if not profiles:
        return lines
    lines += ["CHARACTER REFERENCE SHEETS", "=" * 26, ""]
    for profile in profiles:
        if not isinstance(profile, dict):
            continue
        name = _safe(profile.get("character_name"))
        if not name:
            continue
        role = _profile_role(profile)
        header = name + (f" - {role}" if role else "")
        lines += [header, "-" * len(header)]
        details = _appearance_block(profile)
        for key, label in [
            ("age_range", "Age"),
            ("gender_presentation", "Gender presentation"),
            ("height", "Height"),
            ("body_type", "Body type"),
            ("silhouette_shape", "Silhouette"),
            ("face_shape", "Face shape"),
            ("skin_tone_or_markings", "Skin / markings"),
            ("hair_style", "Hair style"),
            ("hair_color", "Hair color"),
            ("eye_shape", "Eye shape"),
            ("eye_color", "Eye color"),
            ("clothing_style", "Clothing style"),
            ("main_outfit_description", "Main outfit"),
            ("iconic_item", "Iconic item"),
            ("expression_style", "Expression style"),
            ("pose_language", "Pose language"),
            ("manga_panel_presence", "Panel presence"),
            ("first_impression_visual", "First impression"),
        ]:
            v = details.get(key)
            text = ", ".join(_text_list(v)) if isinstance(v, list) else _safe(v)
            if text:
                lines.append(f"  {label}: {text}")
        for key, label in [
            ("distinctive_features", "Distinctive features"),
            ("scars_or_birthmarks", "Scars / birthmarks"),
            ("alternate_outfits", "Alternate outfits"),
            ("accessories", "Accessories"),
            ("weapons_or_tools_visible", "Weapons / tools"),
            ("color_palette", "Color palette"),
        ]:
            text = ", ".join(_text_list(details.get(key, [])))
            if text:
                lines.append(f"  {label}: {text}")
        symbol = _safe(details.get("visual_symbol_or_motif"))
        if symbol:
            lines.append(f"  Visual symbol / motif: {symbol}")
        contrast = _safe(details.get("visual_contrast_with_other_characters"))
        if contrast:
            lines.append(f"  Visual contrast: {contrast}")
        ai_pos = _safe(details.get("ai_image_prompt_notes"))
        ai_neg = _safe(details.get("negative_prompt_notes"))
        if ai_pos:
            lines.append(f"  AI prompt (positive): {ai_pos}")
        if ai_neg:
            lines.append(f"  AI prompt (negative): {ai_neg}")
        lines.append("")
    return lines


def _location_index_lines(scripts: list[dict]) -> list[str]:
    """Cross-chapter index of every location and which panels reference it."""
    locs: dict[str, list[str]] = {}
    for cs in scripts:
        ch_meta = cs.get("chapter_metadata", {}) or {}
        ch_num = _safe(ch_meta.get("chapter_number") or ch_meta.get("chapter_id"), "?")
        scene_breakdown = _as_list(cs.get("chapter_scene_breakdown"))
        scene_locations = {
            _safe(s.get("scene_id")): _safe(s.get("location"))
            for s in scene_breakdown
            if isinstance(s, dict)
        }
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            scene_id = _safe(page.get("scene_id"))
            location = scene_locations.get(scene_id) or _safe(page.get("location"))
            if not location:
                continue
            page_num = _safe(page.get("page_number") or page.get("page_id"))
            ref = f"Ch.{ch_num} Pg.{page_num}"
            locs.setdefault(location, []).append(ref)
    if not locs:
        return []
    lines = ["LOCATION INDEX", "=" * 14, ""]
    for loc in sorted(locs.keys()):
        lines.append(f"{loc}")
        lines.append(f"  Used in: {', '.join(locs[loc])}")
    lines.append("")
    return lines


def _panel_full_block(panel: dict, indent: str = "    ") -> list[str]:
    """Render every production-relevant panel field for the artist."""
    out: list[str] = []
    panel_num = _safe(panel.get("panel_number") or panel.get("panel_id"))
    shot_block = panel.get("camera_shot")
    camera = _selected(shot_block) if isinstance(shot_block, dict) else _safe(shot_block)
    size_block = panel.get("panel_size")
    size = _selected(size_block) if isinstance(size_block, dict) else _safe(size_block)
    pacing_block = panel.get("pacing")
    pacing = _selected(pacing_block) if isinstance(pacing_block, dict) else _safe(pacing_block)
    tags = [t for t in [size, camera, pacing] if t]
    header = f"{indent}Panel {panel_num}"
    if tags:
        header += f" [{' / '.join(tags)}]"
    out.append(header)
    for key, label in [
        ("visual", "Visual"),
        ("character_action", "Action"),
        ("background_details", "Background"),
        ("facial_expression", "Expression"),
        ("pose_or_body_language", "Pose"),
        ("mood", "Mood"),
        ("narration", "Narration"),
        ("continuity_notes", "Continuity"),
        ("custom_panel_details", "Custom"),
    ]:
        v = _safe(panel.get(key))
        if v:
            out.append(f"{indent}  {label}: {v}")
    sfx_items = _as_list(panel.get("sound_effects"))
    sfx_parts = []
    for s in sfx_items:
        if isinstance(s, dict):
            t = _safe(s.get("sfx_text"))
            meaning = _safe(s.get("sfx_meaning"))
            if t:
                sfx_parts.append(f"{t}{f' ({meaning})' if meaning else ''}")
        elif isinstance(s, str) and s.strip():
            sfx_parts.append(s.strip())
    if sfx_parts:
        out.append(f"{indent}  SFX: {', '.join(sfx_parts)}")
    return out


def _assemble_visuals_lines(files: dict, all_scripts: list[dict] | None = None) -> list[str]:
    """Build the visual reference document.

    Aggregates panel content across ALL chapter_script snapshots (one per chapter)
    plus a character reference sheet up top and a location index.
    """
    lines: list[str] = []
    title = _story_title(files.get("master_story", {}) or {}, "Story")
    lines += [f"VISUAL REFERENCE - {title}", "=" * (19 + len(title)), ""]

    # Character reference sheets (Phase 2)
    lines.extend(_character_reference_lines(files.get("characters", {}) or {}))

    # Choose source: cross-version aggregate or single current chapter_script
    scripts: list[dict]
    if all_scripts is not None:
        scripts = all_scripts
    else:
        cs = files.get("chapter_script", {}) or {}
        scripts = [cs] if _as_list(cs.get("pages")) else []

    if not scripts:
        lines.append("No script pages found. Generate the Manga Script for at least one chapter first.")
        return lines

    # Location index (Phase 2)
    lines.extend(_location_index_lines(scripts))

    # Per-chapter panel breakdown (Phases 1 + 3)
    lines += ["CHAPTERS", "=" * 8, ""]
    for cs in scripts:
        metadata = cs.get("chapter_metadata", {}) or {}
        ch_num = _safe(metadata.get("chapter_number") or metadata.get("chapter_id"), "?")
        ch_title = _safe(metadata.get("chapter_title"))
        ch_status = _safe(metadata.get("chapter_status"))
        ch_header = f"Chapter {ch_num}{': ' + ch_title if ch_title else ''}"
        lines += [ch_header, "-" * len(ch_header)]
        if ch_status:
            lines.append(f"Status: {ch_status}")
        scene_breakdown = _as_list(cs.get("chapter_scene_breakdown"))
        scene_titles = {
            _safe(s.get("scene_id")): _safe(s.get("scene_title") or s.get("location"))
            for s in scene_breakdown
            if isinstance(s, dict)
        }
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            page_num = _safe(page.get("page_number") or page.get("page_id"))
            scene_id = _safe(page.get("scene_id"))
            scene_label = scene_titles.get(scene_id, scene_id)
            page_purpose = _safe(page.get("page_purpose"))
            page_mood = _safe(page.get("page_mood"))
            header = f"\n  Page {page_num}"
            if scene_label:
                header += f" - Scene: {scene_label}"
            lines.append(header)
            if page_purpose:
                lines.append(f"    Purpose: {page_purpose}")
            if page_mood:
                lines.append(f"    Mood: {page_mood}")
            for panel in _as_list(page.get("panels")):
                if isinstance(panel, dict):
                    lines.extend(_panel_full_block(panel, indent="    "))
        lines.append("")

    return lines


def _panels_csv(scripts: list[dict]) -> str:
    """Per-panel CSV for spreadsheet workflows. One row per panel across all chapters."""
    import csv
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([
        "chapter_number", "chapter_id", "chapter_title",
        "page_number", "page_id", "scene_id",
        "panel_number", "panel_id", "panel_size", "camera_shot", "pacing",
        "visual", "character_action", "background_details",
        "facial_expression", "pose_or_body_language", "mood",
        "narration", "dialogue_count", "sfx_text",
        "continuity_notes",
    ])
    for cs in scripts:
        meta = cs.get("chapter_metadata", {}) or {}
        ch_num = _safe(meta.get("chapter_number"))
        ch_id = _safe(meta.get("chapter_id"))
        ch_title = _safe(meta.get("chapter_title"))
        for page in _as_list(cs.get("pages")):
            if not isinstance(page, dict):
                continue
            page_num = _safe(page.get("page_number"))
            page_id = _safe(page.get("page_id"))
            scene_id = _safe(page.get("scene_id"))
            for panel in _as_list(page.get("panels")):
                if not isinstance(panel, dict):
                    continue
                shot_block = panel.get("camera_shot")
                size_block = panel.get("panel_size")
                pacing_block = panel.get("pacing")
                sfx_items = _as_list(panel.get("sound_effects"))
                sfx_text = "; ".join(
                    _safe(s.get("sfx_text")) if isinstance(s, dict) else _safe(s)
                    for s in sfx_items
                    if (isinstance(s, dict) and _safe(s.get("sfx_text"))) or (isinstance(s, str) and s.strip())
                )
                writer.writerow([
                    ch_num, ch_id, ch_title,
                    page_num, page_id, scene_id,
                    _safe(panel.get("panel_number")),
                    _safe(panel.get("panel_id")),
                    _selected(size_block) if isinstance(size_block, dict) else _safe(size_block),
                    _selected(shot_block) if isinstance(shot_block, dict) else _safe(shot_block),
                    _selected(pacing_block) if isinstance(pacing_block, dict) else _safe(pacing_block),
                    _safe(panel.get("visual")),
                    _safe(panel.get("character_action")),
                    _safe(panel.get("background_details")),
                    _safe(panel.get("facial_expression")),
                    _safe(panel.get("pose_or_body_language")),
                    _safe(panel.get("mood")),
                    _safe(panel.get("narration")),
                    len(_as_list(panel.get("dialogue"))),
                    sfx_text,
                    _safe(panel.get("continuity_notes")),
                ])
    return buf.getvalue()


def _ai_prompt_files(characters: dict) -> dict[str, str]:
    """One file per character holding their AI image prompt notes."""
    out: dict[str, str] = {}
    profiles = (characters.get("created_major_character_profiles", []) or []) + (characters.get("created_side_character_profiles", []) or [])
    for p in profiles:
        if not isinstance(p, dict):
            continue
        name = _safe(p.get("character_name"))
        if not name:
            continue
        details = _appearance_block(p)
        pos = _safe(details.get("ai_image_prompt_notes"))
        neg = _safe(details.get("negative_prompt_notes"))
        if not pos and not neg:
            continue
        safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in name).strip().replace(" ", "_") or "character"
        body = []
        if pos:
            body.append("# Positive prompt\n" + pos)
        if neg:
            body.append("\n# Negative prompt\n" + neg)
        out[f"{safe_name}.txt"] = "\n".join(body)
    return out


def _character_sheet_files(characters: dict) -> dict[str, str]:
    """One Markdown file per character with their full visual sheet."""
    out: dict[str, str] = {}
    profiles = (characters.get("created_major_character_profiles", []) or []) + (characters.get("created_side_character_profiles", []) or [])
    for p in profiles:
        if not isinstance(p, dict):
            continue
        name = _safe(p.get("character_name"))
        if not name:
            continue
        details = _appearance_block(p)
        if not details:
            continue
        safe_name = "".join(c if c.isalnum() or c in "-_ " else "_" for c in name).strip().replace(" ", "_") or "character"
        md: list[str] = [f"# {name}"]
        role = _profile_role(p)
        if role:
            md.append(f"_{role}_\n")
        for key, label in [
            ("age_range", "Age"),
            ("body_type", "Body type"),
            ("face_shape", "Face shape"),
            ("hair_style", "Hair"),
            ("hair_color", "Hair color"),
            ("eye_color", "Eye color"),
            ("clothing_style", "Clothing"),
            ("main_outfit_description", "Outfit"),
            ("iconic_item", "Iconic item"),
            ("expression_style", "Expression style"),
        ]:
            v = _safe(details.get(key))
            if v:
                md.append(f"**{label}**: {v}")
        for key, label in [
            ("distinctive_features", "Distinctive features"),
            ("accessories", "Accessories"),
            ("weapons_or_tools_visible", "Weapons / tools"),
            ("color_palette", "Color palette"),
        ]:
            text = ", ".join(_text_list(details.get(key, [])))
            if text:
                md.append(f"**{label}**: {text}")
        ai_pos = _safe(details.get("ai_image_prompt_notes"))
        ai_neg = _safe(details.get("negative_prompt_notes"))
        if ai_pos:
            md.append(f"\n## AI prompt (positive)\n{ai_pos}")
        if ai_neg:
            md.append(f"\n## AI prompt (negative)\n{ai_neg}")
        out[f"{safe_name}.md"] = "\n\n".join(md) + "\n"
    return out


# ─── routes ──────────────────────────────────────────────────────────────────

def _profile_role(profile: dict) -> str:
    """Resolve a character profile's role from the canonical template path."""
    role_block = profile.get("character_role_level")
    if isinstance(role_block, dict):
        sel = _safe(role_block.get("selected"))
        if sel:
            return sel
        custom = _safe(role_block.get("custom_character_role_level"))
        if custom:
            return custom
    return _safe(profile.get("profile_label"))


def _profile_bio(profile: dict) -> str:
    """Build a one-paragraph bio from real character template fields."""
    bits: list[str] = []
    backstory = profile.get("character_backstory_mental_state_and_community_place", {}) or {}
    bd = backstory.get("backstory_details", {}) if isinstance(backstory, dict) else {}
    if isinstance(bd, dict):
        for key in ("childhood_summary", "important_past_event", "past_trauma_or_wound"):
            v = _safe(bd.get(key))
            if v:
                bits.append(v)
                break
    personality = profile.get("character_personality", {}) or {}
    pd = personality.get("personality_details", {}) if isinstance(personality, dict) else {}
    if isinstance(pd, dict):
        traits = _text_list(pd.get("core_traits") or pd.get("positive_traits") or [])
        if traits:
            bits.append("Traits: " + ", ".join(traits[:5]))
    powers = profile.get("optional_powers_and_power_level") or profile.get("powers_and_abilities", {}) or {}
    if isinstance(powers, dict) and powers.get("is_enabled"):
        pdetails = powers.get("power_details", {}) if isinstance(powers.get("power_details"), dict) else {}
        pname = _safe(pdetails.get("power_name"))
        if pname:
            bits.append(f"Power: {pname}")
    return " - ".join(bits)


def _world_rules_lines(ms: dict) -> list[str]:
    """Pull selected world rules + their detail blocks from master_story."""
    block = ms.get("world_master_rules", {}) or {}
    if not isinstance(block, dict):
        return []
    selected = _text_list(block.get("selected", []))
    selected = [s for s in selected if s != "Custom"]
    rule_details = block.get("rule_details", {}) if isinstance(block.get("rule_details"), dict) else {}
    out: list[str] = []
    if selected:
        out.append(f"Selected: {', '.join(selected)}")
    detail_labels = [
        ("magic_rules", "Magic"),
        ("power_rules", "Powers"),
        ("demon_rules", "Demons"),
        ("monster_rules", "Monsters"),
        ("god_rules", "Gods"),
        ("technology_rules", "Technology"),
        ("race_species_rules", "Races / Species"),
        ("realm_dimension_rules", "Realms / Dimensions"),
        ("forbidden_rules", "Forbidden"),
        ("power_limits", "Power limits"),
    ]
    for key, label in detail_labels:
        text = _safe(rule_details.get(key))
        if text:
            out.append(f"{label}: {text}")
    return out


def _threats_lines(ms: dict) -> list[str]:
    """Full threat block: major, minor, sources, goals, stakes, time limit."""
    block = ms.get("major_threats_and_minor_side_threats", {}) or {}
    if not isinstance(block, dict):
        return []
    out: list[str] = []
    major = _safe(block.get("major_threat"))
    if major and major != "Custom":
        out.append(f"Major threat: {major}")
    minor = _text_list(block.get("minor_side_threats", []))
    minor = [m for m in minor if m != "Custom"]
    if minor:
        out.append(f"Minor threats: {', '.join(minor)}")
    details = block.get("threat_details", {}) if isinstance(block.get("threat_details"), dict) else {}
    for key, label in [
        ("main_threat_source", "Source"),
        ("main_threat_goal", "Goal"),
        ("main_threat_target", "Target"),
        ("stakes_if_major_threat_wins", "Stakes if it wins"),
        ("threat_level", "Level"),
        ("time_limit", "Time limit"),
        ("hidden_truth_behind_threat", "Hidden truth"),
    ]:
        v = _safe(details.get(key))
        if v:
            out.append(f"{label}: {v}")
    return out


def _relationships_lines(ch: dict) -> list[str]:
    rel_map = ch.get("character_relationship_map", {}) or {}
    if not isinstance(rel_map, dict):
        return []
    rels = _as_list(rel_map.get("relationships", []))
    out: list[str] = []
    for r in rels:
        if not isinstance(r, dict):
            continue
        pair = _safe(r.get("characters_involved"))
        rtype = _safe(r.get("relationship_change_type"))
        reason = _safe(r.get("reason"))
        if not pair:
            continue
        line = pair
        if rtype:
            line += f" - {rtype}"
        if reason:
            line += f" ({reason})"
        out.append(line)
    return out


def _plot_threads_lines(po: dict) -> list[str]:
    threads = po.get("plot_threads", {}) or {}
    if not isinstance(threads, dict):
        return []
    out: list[str] = []
    main = threads.get("main_plot_thread", {}) if isinstance(threads.get("main_plot_thread"), dict) else {}
    main_goal = _safe(main.get("goal"))
    if main_goal:
        out.append(f"Main goal: {main_goal}")
        for key, label in [("obstacles", "Obstacles"), ("turning_points", "Turning points"), ("resolution", "Resolution")]:
            v = main.get(key)
            text = ", ".join(_text_list(v)) if isinstance(v, list) else _safe(v)
            if text:
                out.append(f"  {label}: {text}")
    arcs = _as_list(threads.get("character_arc_threads", []))
    for arc in arcs:
        if not isinstance(arc, dict):
            continue
        cid = _safe(arc.get("character_id"))
        start = _safe(arc.get("starting_state"))
        final = _safe(arc.get("final_state"))
        if cid or start or final:
            out.append(f"Character arc {cid}: {start or '?'} -> {final or '?'}")
    return out


def _assemble_story_lines(files: dict) -> list[str]:
    """Build a clean readable story document from the current official files."""
    lines: list[str] = []
    ms = files.get("master_story", {})
    ch = files.get("characters", {})
    po = files.get("plot_outline", {})
    cs = files.get("chapter_script", {})

    title = _story_title(ms)
    lines += [title, "=" * len(title), ""]

    # Subtitle: story types + foundation + ending direction
    story_types = _selected(ms.get("story_type"))
    foundation = _selected(ms.get("story_foundation"))
    ending = _selected(ms.get("ending_direction"))
    subtitle_bits = [b for b in [story_types, foundation, ending] if b]
    if subtitle_bits:
        lines += [" | ".join(subtitle_bits), ""]

    # SYNOPSIS — idea_so_far is the canonical premise field in the template
    premise = _safe(ms.get("idea_so_far"))
    if premise:
        lines += ["SYNOPSIS", "-" * 8, premise, ""]

    # WORLD — world_type + selected world rules + their detail blocks
    world_type = _selected(ms.get("world_type"))
    rules_lines = _world_rules_lines(ms)
    if world_type or rules_lines:
        lines += ["WORLD", "-" * 5]
        _append_field(lines, "Type", world_type)
        lines.extend(rules_lines)
        lines.append("")

    # FACTIONS — selected list
    factions_block = ms.get("major_factions_and_ruling_sides", {}) or {}
    factions = _text_list(factions_block.get("selected", []) if isinstance(factions_block, dict) else [])
    factions = [f for f in factions if f != "Custom"]
    if factions:
        lines += ["FACTIONS", "-" * 8, ", ".join(factions), ""]

    # THREATS — full block (major + minor + source/goal/stakes/level)
    threat_lines = _threats_lines(ms)
    if threat_lines:
        lines += ["THREATS", "-" * 7]
        lines.extend(threat_lines)
        lines.append("")

    # CHARACTERS — real role + bio from canonical fields
    profiles = (ch.get("created_major_character_profiles", []) or []) + (ch.get("created_side_character_profiles", []) or [])
    if profiles:
        lines += ["CHARACTERS", "-" * 10]
        for profile in profiles:
            if not isinstance(profile, dict):
                continue
            name = _safe(profile.get("character_name") or profile.get("name"))
            if not name:
                continue
            role = _profile_role(profile)
            bio = _profile_bio(profile)
            lines.append(f"{name}{' - ' + role if role else ''}")
            if bio:
                lines.append(f"  {bio}")
        lines.append("")

    # RELATIONSHIPS — character_relationship_map
    rel_lines = _relationships_lines(ch)
    if rel_lines:
        lines += ["RELATIONSHIPS", "-" * 13]
        lines.extend(rel_lines)
        lines.append("")

    # ARC OVERVIEW + PLOT THREADS
    story_arc = po.get("story_arc_overview", {}) or {}
    narrative_structure = _selected(po.get("narrative_structure"))
    arc_title = _safe(story_arc.get("arc_title"), "Current Arc")
    lines += ["ARC OVERVIEW", "-" * 12]
    _append_field(lines, "Structure", narrative_structure)
    _append_field(lines, "Arc", arc_title)
    _append_field(lines, "Arc length", _selected(story_arc.get("arc_length_type")))
    for label, key in [
        ("Summary", "arc_summary"),
        ("Story question", "main_story_question"),
        ("Emotional question", "central_emotional_question"),
        ("External conflict", "main_external_conflict"),
        ("Internal conflict", "main_internal_conflict"),
        ("Relationship conflict", "main_relationship_conflict"),
        ("Ending target", "ending_type_target"),
    ]:
        _append_field(lines, label, story_arc.get(key))
    lines.append("")

    thread_lines = _plot_threads_lines(po)
    if thread_lines:
        lines += ["PLOT THREADS", "-" * 12]
        lines.extend(thread_lines)
        lines.append("")

    # CHAPTERS + per-chapter script content
    outline_chapters = _as_list((po.get("chapter_or_episode_list", {}) or {}).get("chapters", []))
    scene_cards = _as_list((po.get("scene_cards", {}) or {}).get("scenes", []))
    scenes_by_chapter = _group_scenes_by_chapter(scene_cards)
    script_pages = _as_list(cs.get("pages"))
    script_chapter_id = _safe((cs.get("chapter_metadata") or {}).get("chapter_id"))

    if not outline_chapters:
        lines += ["CHAPTERS", "=" * 8, "", "No chapters found. Create chapters on the Plot Board first.", ""]
        return lines

    lines += ["FULL STORY", "=" * 10, ""]
    for chapter in sorted(outline_chapters, key=lambda c: int(c.get("chapter_number") or 0) if isinstance(c, dict) else 0):
        if not isinstance(chapter, dict):
            continue
        header = _chapter_header(chapter)
        chapter_id = _safe(chapter.get("chapter_id"))
        chapter_scenes = scenes_by_chapter.get(chapter_id, [])
        # The chapter_script.json holds ONE chapter at a time. Match against its chapter_metadata.chapter_id.
        chapter_script_pages = script_pages if (chapter_id and chapter_id == script_chapter_id) else []

        lines += [header, "-" * len(header)]
        _append_field(lines, "Arc", chapter.get("arc_title") or arc_title)
        _append_field(lines, "Structure beat", chapter.get("structure_section"))
        lines.append("")

        lines += ["Story", "-----"]
        script_lines = _panel_story_lines(chapter_script_pages)
        if script_lines:
            lines.extend(script_lines)
        else:
            for key in [
                "chapter_purpose",
                "summary",
                "main_conflict",
                "emotional_beat",
                "twist_or_hook",
                "ending_cliffhanger",
                "custom_chapter_details",
            ]:
                text = _safe(chapter.get(key))
                if text:
                    lines.append(text)
            if chapter_scenes:
                lines.append("")
                lines.append("Key scenes")
                seen_scene_beats: set[str] = set()
                for scene in chapter_scenes:
                    scene_order = _safe(scene.get("scene_order") or scene.get("scene_id"), "?")
                    scene_goal = _safe(scene.get("scene_goal"))
                    scene_reveal = _safe(scene.get("new_information_revealed"))
                    scene_ending = _safe(scene.get("ending_beat"))
                    scene_bits = [scene_goal]
                    if scene_reveal and scene_reveal != scene_goal:
                        scene_bits.append(scene_reveal)
                    if scene_ending and scene_ending not in scene_bits:
                        scene_bits.append(scene_ending)
                    scene_text = " ".join(bit for bit in scene_bits if bit)
                    if scene_text in seen_scene_beats:
                        continue
                    seen_scene_beats.add(scene_text)
                    if scene_text:
                        lines.append(f"{scene_order}. {scene_text}")
        lines.append("")

        lines += ["Chapter Review", "--------------"]
        _append_field(lines, "Characters active", chapter.get("characters_present"))
        _append_field(lines, "Relationship movement", chapter.get("relationships_used"))
        _append_field(lines, "Threat / faction movement", _text_list(chapter.get("threats_used")) + _text_list(chapter.get("factions_used")))
        _append_field(lines, "World or powers", _text_list(chapter.get("world_rules_shown")) + _text_list(chapter.get("power_system_shown")))
        _append_field(lines, "Open hook", chapter.get("ending_cliffhanger") or chapter.get("twist_or_hook"))
        lines.append("")

    return lines


def _script_pages_by_scene(cs: dict) -> dict[str, list[dict]]:
    """Group chapter_script pages by scene_id."""
    grouped: dict[str, list[dict]] = {}
    for page in _as_list(cs.get("pages", [])):
        if not isinstance(page, dict):
            continue
        sid = _safe(page.get("scene_id"))
        if sid:
            grouped.setdefault(sid, []).append(page)
    return grouped


def _scene_script_lines(pages: list[dict], indent: str = "  ") -> list[str]:
    """Render the script content for one scene: per-panel narration, dialogue, SFX, visuals."""
    out: list[str] = []
    for page in pages:
        page_label = _safe(page.get("page_number") or page.get("page_id"))
        out.append(f"{indent}- Page {page_label}")
        for panel in _as_list(page.get("panels", [])):
            if not isinstance(panel, dict):
                continue
            panel_no = _safe(panel.get("panel_number") or panel.get("panel_id"))
            shot_block = panel.get("camera_shot")
            shot = _selected(shot_block) if isinstance(shot_block, dict) else _safe(shot_block)
            visual = _safe(panel.get("visual"))
            action = _safe(panel.get("character_action"))
            narration = _safe(panel.get("narration"))
            header = f"{indent}  Panel {panel_no}"
            if shot:
                header += f" [{shot}]"
            out.append(header)
            if visual:
                out.append(f"{indent}    Visual: {visual}")
            if action:
                out.append(f"{indent}    Action: {action}")
            if narration:
                out.append(f"{indent}    Narration: {narration}")
            for d in _as_list(panel.get("dialogue", [])):
                if not isinstance(d, dict):
                    continue
                text = _safe(d.get("text"))
                if not text:
                    continue
                speaker = _safe(d.get("speaker_name") or d.get("speaker") or d.get("speaker_id")) or "?"
                bubble_block = d.get("speech_bubble_type")
                bubble = _selected(bubble_block) if isinstance(bubble_block, dict) else _safe(bubble_block)
                bubble_tag = f" ({bubble})" if bubble and bubble != "Normal" else ""
                out.append(f'{indent}    {speaker}{bubble_tag}: "{text}"')
            sfx_items = _as_list(panel.get("sound_effects", []))
            sfx_texts = []
            for s in sfx_items:
                if isinstance(s, dict):
                    t = _safe(s.get("sfx_text"))
                    if t:
                        sfx_texts.append(t)
                elif isinstance(s, str) and s.strip():
                    sfx_texts.append(s.strip())
            if sfx_texts:
                out.append(f"{indent}    SFX: {', '.join(sfx_texts)}")
    return out


def _assemble_scenes_lines(files: dict) -> list[str]:
    """Build full scene-card export with chapter-script dialogue/SFX/visuals injected per scene."""
    po = files.get("plot_outline", {})
    ms = files.get("master_story", {}) or {}
    cs = files.get("chapter_script", {}) or {}
    lines: list[str] = []

    title = _story_title(ms, "Story")
    lines += [f"SCENES - {title}", "=" * (9 + len(title)), ""]

    chapters = _as_list((po.get("chapter_or_episode_list", {}) or {}).get("chapters", []))
    scene_cards = _as_list((po.get("scene_cards", {}) or {}).get("scenes", []))
    scenes_by_chapter = _group_scenes_by_chapter(scene_cards)
    pages_by_scene = _script_pages_by_scene(cs)
    script_chapter_id = _safe((cs.get("chapter_metadata") or {}).get("chapter_id"))

    if not chapters and not scene_cards:
        lines.append("No scenes found. Add scenes in the Scene Cards studio.")
        return lines

    chapter_ids_seen: set[str] = set()
    for chapter in sorted(chapters, key=lambda c: int(c.get("chapter_number") or 0) if isinstance(c, dict) else 0):
        if not isinstance(chapter, dict):
            continue
        chapter_id = _safe(chapter.get("chapter_id"))
        chapter_ids_seen.add(chapter_id)
        header = _chapter_header(chapter)
        lines += [header, "-" * len(header)]
        _append_field(lines, "Chapter purpose", chapter.get("chapter_purpose"))
        _append_field(lines, "Chapter summary", chapter.get("summary"))
        _append_field(lines, "Chapter conflict", chapter.get("main_conflict"))
        _append_field(lines, "Chapter hook", chapter.get("ending_cliffhanger") or chapter.get("twist_or_hook"))
        if chapter_id and chapter_id == script_chapter_id:
            lines.append("Script: AVAILABLE (panel content shown per scene below)")
        elif chapter_id:
            lines.append("Script: not generated for this chapter yet")

        chapter_scenes = scenes_by_chapter.get(chapter_id, [])
        if not chapter_scenes:
            lines.append("No scene cards for this chapter.")
            lines.append("")
            continue

        lines.append("")
        for scene in chapter_scenes:
            scene_no = _safe(scene.get("scene_order") or scene.get("scene_id"), "?")
            scene_title = _safe(scene.get("scene_title") or scene.get("scene_id"), f"Scene {scene_no}")
            scene_id = _safe(scene.get("scene_id"))
            lines.append(f"Scene {scene_no}: {scene_title}")
            _append_field(lines, "Location", scene.get("location"), "  ")
            _append_field(lines, "Time", scene.get("time"), "  ")
            _append_field(lines, "Characters", scene.get("characters_present"), "  ")
            _append_field(lines, "Goal", scene.get("scene_goal"), "  ")
            _append_field(lines, "Conflict", scene.get("scene_conflict"), "  ")
            _append_field(lines, "Relationship dynamic", scene.get("relationship_dynamic_used"), "  ")
            _append_field(lines, "New information", scene.get("new_information_revealed"), "  ")
            _append_field(lines, "Action/dialogue focus", scene.get("action_or_dialogue_focus"), "  ")
            _append_field(lines, "Visual manga moment", scene.get("visual_manga_moment"), "  ")
            _append_field(lines, "Panel mood", scene.get("panel_mood"), "  ")
            _append_field(lines, "Ending beat", scene.get("ending_beat"), "  ")
            _append_field(lines, "Custom details", scene.get("custom_scene_details"), "  ")

            # Inject script content if this scene has generated pages.
            scene_pages = pages_by_scene.get(scene_id, []) if scene_id else []
            if scene_pages:
                lines.append("  Script:")
                lines.extend(_scene_script_lines(scene_pages, indent="  "))
            lines.append("")

    orphan_scenes = [scene for scene in scene_cards if isinstance(scene, dict) and _safe(scene.get("chapter_id")) not in chapter_ids_seen]
    if orphan_scenes:
        lines += ["Unassigned Scenes", "-" * 17]
        for scene in orphan_scenes:
            lines.append(_safe(scene.get("scene_id"), "Scene"))
            _append_field(lines, "Chapter ID", scene.get("chapter_id"), "  ")
            _append_field(lines, "Goal", scene.get("scene_goal"), "  ")
            _append_field(lines, "Visual manga moment", scene.get("visual_manga_moment"), "  ")
            lines.append("")

    return lines


@router.get("/story")
def export_story(
    story_id: str,
    fmt: Literal["txt", "md", "docx"] = Query(default="md"),
    registry: SQLiteRegistry = Depends(get_registry),
):
    files = _get_all_files(story_id, registry)
    title = _safe(
        (files.get("master_story", {}) or {}).get("title")
        or (files.get("master_story", {}) or {}).get("story_title"),
        "story",
    )
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip().replace(" ", "_")
    lines = _assemble_story_lines(files)

    if fmt == "txt":
        content = _lines_to_text(lines).encode("utf-8")
        media = "text/plain; charset=utf-8"
        filename = f"{safe_title}.txt"
    elif fmt == "md":
        content = _lines_to_markdown(lines).encode("utf-8")
        media = "text/markdown; charset=utf-8"
        filename = f"{safe_title}.md"
    else:  # docx
        content = _lines_to_docx(lines, title)
        media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{safe_title}.docx"

    return Response(
        content=content,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/scenes")
def export_scenes(
    story_id: str,
    fmt: Literal["txt", "md"] = Query(default="md"),
    registry: SQLiteRegistry = Depends(get_registry),
):
    files = _get_all_files(story_id, registry)
    title = _safe(
        (files.get("master_story", {}) or {}).get("title")
        or (files.get("master_story", {}) or {}).get("story_title"),
        "story",
    )
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip().replace(" ", "_")
    lines = _assemble_scenes_lines(files)

    if fmt == "txt":
        content = _lines_to_text(lines).encode("utf-8")
        media = "text/plain; charset=utf-8"
        filename = f"{safe_title}_scenes.txt"
    else:
        content = _lines_to_markdown(lines).encode("utf-8")
        media = "text/markdown; charset=utf-8"
        filename = f"{safe_title}_scenes.md"

    return Response(
        content=content,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/visuals")
def export_visuals(
    story_id: str,
    fmt: Literal["txt", "md"] = Query(default="md"),
    registry: SQLiteRegistry = Depends(get_registry),
):
    files = _get_all_files(story_id, registry)
    title = _safe(
        (files.get("master_story", {}) or {}).get("title")
        or (files.get("master_story", {}) or {}).get("story_title"),
        "story",
    )
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip().replace(" ", "_")
    all_scripts = _all_chapter_scripts(story_id, registry)
    lines = _assemble_visuals_lines(files, all_scripts=all_scripts)

    if fmt == "txt":
        content = _lines_to_text(lines).encode("utf-8")
        media = "text/plain; charset=utf-8"
        filename = f"{safe_title}_visuals.txt"
    else:
        content = _lines_to_markdown(lines).encode("utf-8")
        media = "text/markdown; charset=utf-8"
        filename = f"{safe_title}_visuals.md"

    return Response(
        content=content,
        media_type=media,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/visuals-bundle")
def export_visuals_bundle(
    story_id: str,
    registry: SQLiteRegistry = Depends(get_registry),
):
    """Production-ready ZIP bundle for the artist:
        visuals.md            — full reference document (all chapters)
        panels.csv            — one row per panel for spreadsheet workflows
        character_sheets/*.md — per-character visual sheet
        prompts/*.txt         — per-character AI image prompts (positive + negative)
    """
    files = _get_all_files(story_id, registry)
    characters = files.get("characters", {}) or {}
    title = _safe(
        (files.get("master_story", {}) or {}).get("title")
        or (files.get("master_story", {}) or {}).get("story_title"),
        "story",
    )
    safe_title = "".join(c if c.isalnum() or c in "-_ " else "_" for c in title).strip().replace(" ", "_") or story_id
    all_scripts = _all_chapter_scripts(story_id, registry)

    visuals_md = _lines_to_markdown(_assemble_visuals_lines(files, all_scripts=all_scripts))
    panels_csv = _panels_csv(all_scripts)
    sheet_files = _character_sheet_files(characters)
    prompt_files = _ai_prompt_files(characters)

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("visuals.md", visuals_md)
        zf.writestr("panels.csv", panels_csv)
        readme = (
            f"# {title} — Visual Production Bundle\n\n"
            "- visuals.md: full per-chapter panel breakdown\n"
            "- panels.csv: one row per panel (open in Excel / Sheets)\n"
            "- character_sheets/: per-character visual reference\n"
            "- prompts/: AI image prompts per character (paste into your generator)\n"
        )
        zf.writestr("README.md", readme)
        for name, content in sheet_files.items():
            zf.writestr(f"character_sheets/{name}", content)
        for name, content in prompt_files.items():
            zf.writestr(f"prompts/{name}", content)

    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}_visuals_bundle.zip"'},
    )


@router.get("/raw-zip")
def export_raw_zip(
    story_id: str,
    registry: SQLiteRegistry = Depends(get_registry),
):
    """Stream a ZIP archive containing all current official story JSON files."""
    file_types = ["master_story", "characters", "plot_outline", "memory_system", "plot_workspace", "chapter_script"]
    filename_map = {
        "master_story": "master_story.json",
        "characters": "characters.json",
        "plot_outline": "plot_outline.json",
        "memory_system": "memory_system.json",
        "plot_workspace": "plot_workspace.json",
        "chapter_script": "chapter_script.json",
    }

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, mode="w", compression=zipfile.ZIP_DEFLATED) as zf:
        for ft in file_types:
            rec = registry.get_current_file(story_id, ft)
            data = rec.get("json_copy", {}) if rec else {}
            zf.writestr(filename_map[ft], json.dumps(data, indent=2, ensure_ascii=False))

    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{story_id}_raw_files.zip"'},
    )
